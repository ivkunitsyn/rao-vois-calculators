from __future__ import annotations

import io
from typing import Any

from .formatting import format_internet_component_label, format_money, format_number, format_percent
from .model import ContractTerms, MinimumLicense, RadioReportModel, ReportRow


def build_radio_report_docx(model: RadioReportModel, *, disclaimer: str = "") -> bytes:
    from docx import Document  # type: ignore
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    def set_run(run, *, bold: bool = False, size: float = 10, color: str = "142134", italic: bool = False) -> None:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor.from_string(color)

    def shape_paragraph(
        paragraph,
        *,
        before: float = 0,
        after: float = 4,
        keep_with_next: bool = False,
        keep_together: bool = False,
    ) -> None:
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.keep_with_next = keep_with_next
        fmt.keep_together = keep_together

    def add_heading(text: str, level: int = 2):
        p = doc.add_heading(text, level=level)
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shape_paragraph(p, before=0, after=12, keep_with_next=True)
        else:
            shape_paragraph(p, before=14, after=7, keep_with_next=True)
        for run in p.runs:
            set_run(run, bold=True, size=16 if level == 1 else 12, color="142134" if level == 1 else "204C72")
        return p

    def add_text(
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        size: float = 10,
        color: str = "142134",
        before: float = 0,
        after: float = 4,
        keep_with_next: bool = False,
        keep_together: bool = False,
    ):
        p = doc.add_paragraph()
        shape_paragraph(p, before=before, after=after, keep_with_next=keep_with_next, keep_together=keep_together)
        run = p.add_run(text)
        set_run(run, bold=bold, italic=italic, size=size, color=color)
        return p

    def shade_cell(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def set_cell_margins(cell, top: int = 80, start: int = 110, bottom: int = 80, end: int = 110) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
            node = tc_mar.find(qn(f"w:{name}"))
            if node is None:
                node = OxmlElement(f"w:{name}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def format_cell(cell, *, key: bool = False, result: bool = False) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "EEF8F3" if result else ("F7FAFC" if key else "FFFFFF"))
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_together = True
            for run in paragraph.runs:
                set_run(run, bold=key or result, size=9.2, color="0D513F" if result else "142134")

    def add_rows(rows: list[ReportRow]) -> None:
        clean = [x for x in rows if x.value not in (None, "")]
        if not clean:
            return
        table = doc.add_table(rows=len(clean), cols=2)
        table.style = "Table Grid"
        table.autofit = True
        for idx, row in enumerate(clean):
            left, right = table.cell(idx, 0), table.cell(idx, 1)
            left.text = str(row.label)
            right.text = str(row.value)
            is_result = row.key_result and str(row.value).strip() != "—"
            format_cell(left, key=True, result=is_result)
            format_cell(right, result=is_result)

    def add_hyperlink(text: str, url: str) -> None:
        p = doc.add_paragraph()
        if not url:
            run = p.add_run(text)
            set_run(run, bold=True, color="204C72")
            return
        try:
            r_id = doc.part.relate_to(url, RT.HYPERLINK, is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)
            new_run = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            r_pr.append(OxmlElement("w:b"))
            color_node = OxmlElement("w:color")
            color_node.set(qn("w:val"), "0B7A5A")
            r_pr.append(color_node)
            underline = OxmlElement("w:u")
            underline.set(qn("w:val"), "single")
            r_pr.append(underline)
            new_run.append(r_pr)
            text_node = OxmlElement("w:t")
            text_node.text = text
            new_run.append(text_node)
            hyperlink.append(new_run)
            p._p.append(hyperlink)
        except Exception:
            run = p.add_run(text)
            set_run(run, bold=True, color="0B7A5A")

    def add_license_heading(prefix: str, number: str, url: str, suffix: str) -> None:
        p = doc.add_paragraph()
        shape_paragraph(p, before=8, after=3, keep_with_next=True, keep_together=True)
        run = p.add_run(prefix)
        set_run(run, bold=True, size=10)
        if url:
            try:
                r_id = doc.part.relate_to(url, RT.HYPERLINK, is_external=True)
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), r_id)
                new_run = OxmlElement("w:r")
                r_pr = OxmlElement("w:rPr")
                r_pr.append(OxmlElement("w:b"))
                color_node = OxmlElement("w:color")
                color_node.set(qn("w:val"), "0B7A5A")
                r_pr.append(color_node)
                underline = OxmlElement("w:u")
                underline.set(qn("w:val"), "single")
                r_pr.append(underline)
                new_run.append(r_pr)
                text_node = OxmlElement("w:t")
                text_node.text = number
                new_run.append(text_node)
                hyperlink.append(new_run)
                p._p.append(hyperlink)
            except Exception:
                link_run = p.add_run(number)
                set_run(link_run, bold=True, color="0B7A5A")
        else:
            link_run = p.add_run(number)
            set_run(link_run, bold=True, color="0B7A5A")
        end = p.add_run(suffix)
        set_run(end, bold=True, size=10)

    def add_terms(block: ContractTerms) -> None:
        add_text(block.title, bold=True, color="204C72", before=7, after=2, keep_with_next=True)
        if not block.lines and not block.bullets:
            add_text("—", after=2)
        for line in block.lines:
            add_text(line, after=2, keep_together=True)
        for bullet in block.bullets:
            p = doc.add_paragraph(style="List Bullet")
            shape_paragraph(p, before=0, after=1, keep_together=True)
            run = p.add_run(bullet)
            set_run(run, size=9.5)

    def add_minimum_license(lic: MinimumLicense, trailing_rows: list[ReportRow] | None = None) -> None:
        prefix = "Новая лицензия № " if lic.is_new else "Лицензия № "
        add_license_heading(prefix, lic.license_number, lic.license_url, f" — минимальная сумма: {format_money(lic.minimum)} ₽")
        if float(lic.internet_component or 0.0) > 0:
            label = format_internet_component_label(lic.internet_resources)
            add_text(
                f"{label}: {format_money(lic.internet_component)} ₽.",
                bold=True,
                size=9.5,
                before=0,
                after=4,
                keep_with_next=True,
                keep_together=True,
            )
        coeff = "не применяется" if lic.hours_coefficient is None or abs(float(lic.hours_coefficient) - 1.0) < 1e-9 else str(lic.hours_coefficient).replace(".", ",")
        add_rows(
            [
                ReportRow("Численность населения", format_number(lic.population)),
                ReportRow("Диапазон по численности населения", lic.population_range or "—"),
                ReportRow("Интернет-ресурсы", format_number(lic.internet_resources)),
                ReportRow("Всего часов вещания", format_number(lic.weekly_hours)),
                ReportRow("Коэффициент за объём вещания", coeff),
            ]
            + list(trailing_rows or [])
        )

    add_heading("Результат расчёта", level=1)

    add_heading("1. ИСХОДНЫЕ ДАННЫЕ", level=2)
    add_rows(model.source_data)

    add_heading("2. РАСЧЁТ ПРОЦЕНТНОЙ СТАВКИ", level=2)
    for lic in model.rate_licenses:
        prefix = "Новая лицензия № " if lic.is_new else "Лицензия № "
        add_license_heading(prefix, lic.license_number, lic.license_url, f" — {format_percent(lic.rate)}")
        for ch in lic.channels:
            add_text(
                f"{ch.name} — {format_number(ch.weekly_hours)} ч — {format_percent(ch.rate)}",
                bold=True,
                size=9.5,
                before=0,
                after=2,
                keep_together=True,
            )
            if ch.actual_share_percent is not None:
                add_text(
                    f"Фактическая доля использования: {format_percent(ch.actual_share_percent)}. Ставка: {format_percent(ch.rate)}.",
                    size=9,
                    after=2,
                    keep_together=True,
                )
            for topic in ch.topics:
                p = doc.add_paragraph(style="List Bullet")
                shape_paragraph(p, before=0, after=1, keep_together=True)
                run = p.add_run(f"{topic.name} → категория {topic.category} — {format_percent(topic.rate)}")
                set_run(run, size=9)
    add_rows(
        [
            ReportRow("Процентная ставка по договору", format_percent(model.contract_rate), True),
            ReportRow(
                "Расчётная сумма за квартал",
                f"{format_money(model.quarter_amount, precise=True)} ₽" if model.quarter_amount is not None else "—",
            ),
        ]
    )

    add_heading("3. РАСЧЁТ МИНИМАЛЬНОЙ СУММЫ ЗА КВАРТАЛ", level=2)
    if model.minimum_licenses:
        last_idx = len(model.minimum_licenses) - 1
        for idx, lic in enumerate(model.minimum_licenses):
            add_minimum_license(lic, trailing_rows=model.minimum_rows if idx == last_idx else None)
    else:
        add_rows(model.minimum_rows)

    add_heading("4. УСЛОВИЯ ДОГОВОРА", level=2)
    for block in model.contract_terms:
        add_terms(block)

    add_heading("5. КОММЕНТАРИИ К РАСЧЁТУ", level=2)
    if model.comments:
        for comment in model.comments:
            p = doc.add_paragraph(style="List Bullet")
            shape_paragraph(p, before=0, after=1, keep_together=True)
            run = p.add_run(comment)
            set_run(run, size=9.5)
    else:
        add_text("Дополнительных комментариев нет.")

    if disclaimer:
        add_text(disclaimer, italic=True, size=8, color="53657C", before=10, after=0, keep_together=True)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
