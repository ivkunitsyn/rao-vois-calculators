from pathlib import Path
from typing import Any, List, Optional, Literal
from functools import lru_cache
import csv
import html
import os
import io
import json
import shutil
import glob
import gzip
import subprocess
import threading
import re
import uuid
import time
from datetime import datetime
from zoneinfo import ZoneInfo
import openpyxl
import sqlite3

from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, JSONResponse, Response, FileResponse
from pydantic import BaseModel, Field
from calculators.rao_events.backend.calculator import calculate_event_fee as calculate_rao_event_fee
from calculators.vois_events.backend.calculator import calculate_event_fee as calculate_vois_event_fee
from calculators.rao_radio.backend.calculator import run_calc_capture as run_radio_calc_capture
from calculators.rao_radio.result.renderer_docx import build_radio_report_docx
from calculators.rao_radio.result.renderer_html import extract_model_from_html as extract_radio_report_model_from_html
from calculators.rao_radio.result.renderer_text import render_radio_report_text
from tools.rkn_table_bot.rkn_cleaner import (
    POPULATION_MODE_YANDEX,
    clean_rkn_xlsx,
    build_rkn_sqlite,
    build_inn_name_csv,
    verify_rkn_artifacts,
)

# Pydantic v1/v2 совместимость
try:
    from pydantic import field_validator  # type: ignore
    _V2 = True
except Exception:
    from pydantic import validator as field_validator  # type: ignore
    _V2 = False

BASE_DIR = Path(__file__).resolve().parent
PROJECT_DIR = Path(__file__).resolve().parents[3]
INTERFACE_DIR = PROJECT_DIR / "interface"
TV_DATA_DIR = PROJECT_DIR / "calculators" / "rao_tv" / "data"
VOIS_TV_DATA_DIR = PROJECT_DIR / "calculators" / "vois_tv" / "data"
RKN_BOT_DIR = PROJECT_DIR / "tools" / "rkn_table_bot"
INDEX_HTML = INTERFACE_DIR / "index.html"
FAVICON_SVG = INTERFACE_DIR / "favicon.svg"
FAVICON_ICO = INTERFACE_DIR / "favicon.ico"
FAVICON_PNG = INTERFACE_DIR / "favicon-32x32.png"
APPLE_TOUCH_ICON = INTERFACE_DIR / "apple-touch-icon.png"
RKN_UPLOAD_PASSWORD = os.getenv("RKN_WEB_UPDATE_PASSWORD", "").strip() or "RAOrenewal2843"
RKN_ADMIN_USERS_JSON = os.getenv("RKN_WEB_ADMIN_USERS_JSON", "").strip()
RKN_UPLOAD_POPULATION_MODE = os.getenv("RKN_WEB_POPULATION_MODE", POPULATION_MODE_YANDEX).strip().lower() or POPULATION_MODE_YANDEX
RKN_UPLOAD_ALLOW_EMPTY_STATUS = os.getenv("RKN_WEB_ALLOW_EMPTY_STATUS", "1").strip().lower() in {"1", "true", "yes", "on"}
RKN_UPLOAD_GIT_COMMIT = os.getenv("RKN_WEB_GIT_COMMIT", "1").strip().lower() in {"1", "true", "yes", "on"}
RKN_UPLOAD_GIT_PUSH = os.getenv("RKN_WEB_GIT_PUSH", "1").strip().lower() in {"1", "true", "yes", "on"}
RKN_UPLOAD_GIT_REMOTE = os.getenv("RKN_WEB_GIT_REMOTE", "sonic").strip() or "sonic"
RKN_UPLOAD_GIT_BRANCH = os.getenv("RKN_WEB_GIT_BRANCH", "").strip()
RKN_UPLOAD_GIT_AUTHOR_NAME = os.getenv("RKN_WEB_GIT_AUTHOR_NAME", "").strip()
RKN_UPLOAD_GIT_AUTHOR_EMAIL = os.getenv("RKN_WEB_GIT_AUTHOR_EMAIL", "").strip()
def _resolve_rkn_upload_storage_dir() -> Path:
    raw = str(os.getenv("RKN_UPLOAD_STORAGE_DIR", "") or "").strip()
    if raw:
        return Path(raw).expanduser()
    data_mount = Path("/data")
    try:
        if data_mount.exists() and data_mount.is_dir():
            return data_mount / "rkn_upload"
    except Exception:
        pass
    return PROJECT_DIR / "calculators" / "rao_tv" / "data"


RKN_UPLOAD_STORAGE_DIR = _resolve_rkn_upload_storage_dir()
RKN_UPLOAD_LOG_PATH = Path(
    os.getenv("RKN_UPLOAD_LOG_PATH", str(RKN_UPLOAD_STORAGE_DIR / "_upload_log.jsonl")).strip()
)


def _configured_rkn_artifact_path(env_name: str, fallback: Path) -> Path:
    raw = str(os.getenv(env_name, "") or "").strip()
    if raw:
        return Path(raw).expanduser()
    return fallback


def _rkn_upload_output_paths() -> dict[str, Path]:
    return {
        "xlsx": _configured_rkn_artifact_path(
            "RKN_XLSX_PATH",
            TV_DATA_DIR / "Таблица РКН slim.xlsx",
        ),
        "sqlite": _configured_rkn_artifact_path(
            "RKN_DB_PATH",
            TV_DATA_DIR / "Таблица РКН.sqlite",
        ),
        "inn_csv": _configured_rkn_artifact_path(
            "RKN_INN_CSV_PATH",
            TV_DATA_DIR / "inn_name.csv",
        ),
    }


def _report_html_to_text(raw_html: str) -> str:
    """Plain-text fallback for copy/export clients that cannot consume HTML."""
    text = str(raw_html or "")
    text = re.sub(
        r'(?is)<div class="radioReportRow">\s*<div class="radioReportKey">(.*?)</div>\s*<div class="radioReportValue">(.*?)</div>\s*</div>',
        lambda m: f"{m.group(1)}: {m.group(2)}\n",
        text,
    )
    text = re.sub(
        r'(?is)<div class="radioReportLicenseTitle">(.*?)</div>',
        lambda m: f"{m.group(1)}\n",
        text,
    )
    text = re.sub(
        r'(?is)<div class="radioReportChannelTitle">(.*?)</div>',
        lambda m: f"{m.group(1)}\n",
        text,
    )
    replacements = [
        (r"(?is)</h[1-6]>", "\n"),
        (r"(?is)<br\s*/?>", "\n"),
        (r"(?is)</p>", "\n"),
        (r"(?is)</li>", "\n"),
        (r"(?is)</div>", "\n"),
        (r"(?is)</section>", "\n\n"),
        (r"(?is)<li[^>]*>", "• "),
        (r"(?is)<a\b[^>]*>", ""),
    ]
    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text)
    text = re.sub(r"(?is)<[^>]+>", "", text)
    text = html.unescape(text)
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    compact: List[str] = []
    blank = False
    for line in lines:
        if not line:
            if not blank and compact:
                compact.append("")
            blank = True
            continue
        compact.append(line)
        blank = False
    return "\n".join(compact).strip()


_REPORT_SEP_RE = re.compile(r"^[\s─━—\-_=]{8,}$")
_REPORT_URL_RE = re.compile(r"https?://[^\s<>()]+")
_REPORT_KEY_VALUE_RE = re.compile(r"^([^:\n]{1,90}):\s*(.*)$")


def _report_escape_inline(text: Any) -> str:
    raw = str(text or "")
    parts: list[str] = []
    pos = 0
    for match in _REPORT_URL_RE.finditer(raw):
        before = raw[pos:match.start()]
        url = match.group(0).rstrip(".,;)")
        tail = match.group(0)[len(url):]
        parts.append(html.escape(before))
        label = "РКН" if "rkn.gov.ru" in url.lower() else "ссылка"
        parts.append(
            f'<a class="radioReportLicenseLink" href="{html.escape(url, quote=True)}" '
            f'target="_blank" rel="noopener">{label}</a>'
        )
        parts.append(html.escape(tail))
        pos = match.end()
    parts.append(html.escape(raw[pos:]))
    return "".join(parts)


def _report_is_heading(line: str) -> bool:
    value = str(line or "").strip()
    if not value or len(value) > 90 or ":" in value:
        return False
    letters = [ch for ch in value if ch.isalpha()]
    if not letters:
        return False
    upper = sum(1 for ch in letters if ch.upper() == ch)
    known = {
        "ИСХОДНЫЕ ДАННЫЕ",
        "РАСЧЁТ",
        "РАСЧЕТ",
        "УСЛОВИЯ В ДОГОВОРЕ",
        "УСЛОВИЯ ПО ДОГОВОРУ",
        "УСЛОВИЯ В ДОПОЛНИТЕЛЬНОМ СОГЛАШЕНИИ",
        "УСЛОВИЯ ПО ДОПОЛНИТЕЛЬНОМУ СОГЛАШЕНИЮ",
        "ИТОГО",
        "ПРИМЕЧАНИЯ",
        "ПРИМЕЧАНИЯ И ДОПУЩЕНИЯ",
    }
    return value.upper() in known or (upper / max(1, len(letters)) > 0.8 and len(letters) >= 4)


def _report_clean_lines(text: str) -> list[str]:
    lines: list[str] = []
    for raw_line in str(text or "").splitlines():
        line = raw_line.strip()
        if _REPORT_SEP_RE.match(line):
            if lines and lines[-1] != "":
                lines.append("")
            continue
        if not line:
            if lines and lines[-1] != "":
                lines.append("")
            continue
        lines.append(line)
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    return lines


def _report_group_blocks(lines: list[str]) -> list[list[str]]:
    blocks: list[list[str]] = []
    current: list[str] = []
    for line in lines:
        if line == "":
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(line)
    if current:
        blocks.append(current)
    return blocks


def _plain_report_body_html(lines: list[str]) -> str:
    chunks: list[str] = []
    row_items: list[tuple[str, str]] = []
    list_items: list[str] = []

    def flush_rows() -> None:
        nonlocal row_items
        if not row_items:
            return
        rows = []
        for key, value in row_items:
            rows.append(
                '<div class="radioReportRow">'
                f'<div class="radioReportKey">{_report_escape_inline(key)}</div>'
                f'<div class="radioReportValue">{_report_escape_inline(value)}</div>'
                '</div>'
            )
        chunks.append(f'<div class="radioReportRows">{"".join(rows)}</div>')
        row_items = []

    def flush_list() -> None:
        nonlocal list_items
        if not list_items:
            return
        chunks.append(
            '<ul class="radioReportList">'
            + "".join(f"<li>{_report_escape_inline(item)}</li>" for item in list_items)
            + "</ul>"
        )
        list_items = []

    for line in lines:
        stripped = line.strip()
        bullet = re.match(r"^[•◦\-]\s*(.+)$", stripped)
        license_line = re.match(r"^▪\s*(.+)$", stripped)
        kv = _REPORT_KEY_VALUE_RE.match(stripped)
        if license_line:
            flush_rows()
            flush_list()
            chunks.append(
                f'<div class="radioReportLicenseBlock">'
                f'<div class="radioReportLicenseTitle">{_report_escape_inline(license_line.group(1))}</div>'
                f'</div>'
            )
        elif bullet:
            flush_rows()
            list_items.append(bullet.group(1).strip())
        elif kv and not stripped.startswith(("http://", "https://")):
            flush_list()
            row_items.append((kv.group(1).strip(), kv.group(2).strip()))
        else:
            flush_rows()
            flush_list()
            chunks.append(f"<p>{_report_escape_inline(stripped)}</p>")

    flush_rows()
    flush_list()
    return "".join(chunks)


def _plain_result_to_report_html(text: str, fallback_title: str = "Результат расчёта") -> str:
    """Wrap legacy plain-text calculator output in the same report shell as RAO radio."""
    raw = str(text or "").strip()
    if not raw:
        return ""
    if raw.lstrip().startswith("<article"):
        return raw

    lines = _report_clean_lines(raw)
    if not lines:
        return ""

    title = "Результат расчёта"
    sections: list[tuple[str, list[str]]] = []
    lead: list[str] = []
    current_title: str | None = None
    current_lines: list[str] = []

    def flush_section() -> None:
        nonlocal current_title, current_lines
        if current_title:
            sections.append((current_title, current_lines))
        elif current_lines:
            lead.extend(current_lines)
        current_title = None
        current_lines = []

    first_nonempty = lines[0]
    if not _report_is_heading(first_nonempty):
        lead.append(first_nonempty)
        lines = lines[1:]

    for line in lines:
        if line == "":
            if current_lines and current_lines[-1] != "":
                current_lines.append("")
            elif lead and lead[-1] != "":
                lead.append("")
            continue
        if _report_is_heading(line):
            flush_section()
            current_title = line
            current_lines = []
        else:
            if current_title:
                current_lines.append(line)
            else:
                lead.append(line)
    flush_section()

    article: list[str] = [f'<article class="radioReport"><h1>{_report_escape_inline(title)}</h1>']
    if lead:
        lead_blocks = _report_group_blocks(lead)
        lead_html = "".join(_plain_report_body_html(block) for block in lead_blocks)
        if lead_html:
            article.append(f'<section class="radioReportSection">{lead_html}</section>')
    for section_title, section_lines in sections:
        section_blocks = _report_group_blocks(section_lines)
        body = "".join(_plain_report_body_html(block) for block in section_blocks)
        article.append(
            f'<section class="radioReportSection">'
            f'<h3>{_report_escape_inline(section_title)}</h3>'
            f'{body}'
            f'</section>'
        )
    article.append("</article>")
    return "".join(article)


def _event_fmt_int(value: Any) -> str:
    try:
        return f"{int(round(float(value))):,}".replace(",", " ")
    except Exception:
        return str(value or "0")


def _event_money(value: Any) -> str:
    return f"{_event_fmt_int(value)} ₽"


def _event_music_share_label(value: Any, provider: str) -> str:
    music_object = "фонограмм" if provider == "ВОИС" else "произведений"
    raw = str(value or "100")
    if raw == "unknown":
        return f"доля {music_object} не подтверждена"
    if raw == "100":
        return f"{music_object}: 80% и более"
    return f"{music_object}: до {raw}%"


def _event_mode_label(payload: dict[str, Any]) -> str:
    user_type = str(payload.get("user_type") or "")
    special_type = str(payload.get("special_calculation_type") or "")
    if user_type == "special" and special_type == "quarterly":
        return "Специальная категория 3.2.1"
    if user_type == "special" and special_type == "hourly":
        return "Специальная категория 3.2.2"
    return "Обычный режим"


def _event_access_label(payload: dict[str, Any]) -> str:
    user_type = str(payload.get("user_type") or "")
    special_type = str(payload.get("special_calculation_type") or "")
    if user_type == "special" and special_type == "quarterly":
        return "не учитывается в спецкатегории 3.2.1"
    if user_type == "special" and special_type == "hourly":
        return "не применяется к режиму 3.2.2"
    return "бесплатный" if bool(payload.get("is_free_access")) else "платный"


def _event_kv(label: str, value: Any) -> str:
    return (
        '<div class="radioReportFact">'
        f'<strong>{_report_escape_inline(label)}:</strong> {_report_escape_inline(value)}'
        '</div>'
    )


def _event_line(label: str, value: Any) -> str:
    return (
        '<div class="eventReportLine">'
        f'<strong>{_report_escape_inline(label)}:</strong> '
        f'<span>{_report_escape_inline(value)}</span>'
        '</div>'
    )


def _event_breakdown_line(item: dict[str, Any]) -> str:
    label = str(item.get("label") or "Шаг расчёта")
    formula = str(item.get("formula") or "").strip()
    value = item.get("value")
    result = _event_money(value) if isinstance(value, (int, float)) else str(value or "—")
    compact = f"{formula} = {result}" if formula else result
    return _event_line(label, compact)


def _event_breakdown_summary_line(item: dict[str, Any]) -> str:
    label = str(item.get("label") or "Шаг расчёта")
    value = item.get("value")
    result = _event_money(value) if isinstance(value, (int, float)) else str(value or "—")
    return _event_line(label, result)


def _event_result_to_report_html(
    result: dict[str, Any],
    payload: dict[str, Any],
    *,
    provider: str,
    title: str,
) -> str:
    total = result.get("total", 0)
    breakdown = [x for x in (result.get("breakdown") or []) if isinstance(x, dict)]
    contract_terms = [x for x in (result.get("contract_terms") or []) if isinstance(x, dict)]
    admin_count = int(payload.get("admin_platforms_count") or 0)
    external_count = int(payload.get("external_platforms_count") or 0)
    broadcasts_count = int(payload.get("broadcasts_count") or 1)

    source_rows = [
        ("Режим расчёта", _event_mode_label(payload)),
        ("Доступ к трансляции", _event_access_label(payload)),
        ("Доля", _event_music_share_label(payload.get("music_share"), provider)),
        ("Администрируемые площадки", _event_fmt_int(admin_count)),
        ("Неадминистрируемые площадки", _event_fmt_int(external_count)),
    ]
    if payload.get("region"):
        source_rows.append(("Регион", payload.get("region")))
    if payload.get("visitors") is not None:
        source_rows.append(("Посещения", _event_fmt_int(payload.get("visitors"))))
    if payload.get("duration_hours"):
        source_rows.append(("Длительность трансляции", f"{payload.get('duration_hours'):g} ч."))
    if payload.get("duration_hours_special"):
        source_rows.append(("Длительность трансляции", f"{payload.get('duration_hours_special'):g} ч."))
    if broadcasts_count > 1:
        source_rows.append(("Количество трансляций", _event_fmt_int(broadcasts_count)))

    key_lines = "".join(_event_breakdown_summary_line(item) for item in breakdown)
    detail_lines = "".join(_event_breakdown_line(item) for item in breakdown)
    contract_lines = []
    for item in contract_terms:
        label = str(item.get("label") or "").strip()
        value = str(item.get("value") or "").strip()
        note = str(item.get("note") or "").strip()
        if not label:
            continue
        contract_lines.append(_event_line(label, value or note or "—"))
    if not contract_lines:
        contract_lines.append('<p class="radioReportMuted">Условия договора не сформированы для выбранного режима.</p>')

    return "".join(
        [
            '<article class="radioReport eventReport">',
            '<h1>Результат расчёта</h1>',
            f'<div class="eventReportSubtitle">{_report_escape_inline(title)}</div>',
            f'<div class="eventReportTotal">ИТОГО К ОПЛАТЕ: {_event_money(total)}</div>',
            '<section class="radioReportSection">',
            '<h3>1. ИСХОДНЫЕ ДАННЫЕ</h3>',
            '<div class="radioReportFacts">',
            "".join(_event_kv(label, value) for label, value in source_rows),
            '</div>',
            '</section>',
            '<section class="radioReportSection">',
            '<h3>2. КЛЮЧЕВЫЕ ПАРАМЕТРЫ РАСЧЁТА</h3>',
            '<div class="eventReportCompactList">',
            key_lines or '<p class="radioReportMuted">Нет дополнительных параметров расчёта.</p>',
            '</div>',
            '<details class="eventReportDetails" data-docx-skip="true">',
            '<summary>Показать подробности расчёта</summary>',
            '<div class="eventReportCompactList">',
            detail_lines or '<p class="radioReportMuted">Подробности расчёта отсутствуют.</p>',
            '</div>',
            '</details>',
            '</section>',
            '<section class="radioReportSection">',
            '<h3>3. ИТОГ РАСЧЁТА</h3>',
            '<div class="eventReportCompactList">',
            _event_line("Сумма к оплате", _event_money(total)),
            '</div>',
            '</section>',
            '<section class="radioReportSection">',
            '<h3>4. УСЛОВИЯ ДОГОВОРА</h3>',
            '<div class="eventReportCompactList">',
            "".join(contract_lines),
            '</div>',
            '</section>',
            '</article>',
        ]
    )
RKN_UPLOAD_STATE_PATH = Path(
    os.getenv("RKN_UPLOAD_STATE_PATH", str(RKN_UPLOAD_STORAGE_DIR / "_upload_last.json")).strip()
)
LEGACY_RKN_UPLOAD_LOG_PATH = PROJECT_DIR / "calculators" / "rao_tv" / "data" / "_upload_log.jsonl"
LEGACY_RKN_UPLOAD_STATE_PATH = PROJECT_DIR / "calculators" / "rao_tv" / "data" / "_upload_last.json"
RKN_UPLOAD_LOG_SEED_PATH = PROJECT_DIR / "calculators" / "rao_tv" / "data" / "_upload_log_seed.jsonl"
LEGACY_USAGE_STATS_DB_PATH = PROJECT_DIR / "calculators" / "rao_tv" / "data" / "_usage_stats.sqlite"
LEGACY_USAGE_EVENTS_LOG_PATH = PROJECT_DIR / "calculators" / "rao_tv" / "data" / "_usage_events.jsonl"
USAGE_BACKFILL_LOG_GLOBS = os.getenv(
    "RKN_USAGE_BACKFILL_LOG_GLOBS",
    "/app/logs/*.log;/app/logs/*.log.*;/var/log/nginx/access.log*;/var/log/nginx/*access*.log*;/var/lib/docker/containers/*/*.log",
).strip()
RKN_UPLOAD_GIT_AUTOSYNC = os.getenv("RKN_WEB_GIT_AUTOSYNC", "0").strip().lower() in {"1", "true", "yes", "on"}
RKN_UPLOAD_GIT_AUTOSYNC_SEC = max(5, int(os.getenv("RKN_WEB_GIT_AUTOSYNC_SEC", "20").strip() or "20"))
USAGE_AUTO_RECONCILE_ON_STARTUP = os.getenv("RKN_USAGE_AUTO_RECONCILE_ON_STARTUP", "1").strip().lower() in {"1", "true", "yes", "on"}
_RKN_UPLOAD_LOCK = threading.Lock()
_RKN_UPLOAD_JOBS: dict[str, dict[str, Any]] = {}
_RKN_UPLOAD_JOBS_LOCK = threading.Lock()
_RKN_GIT_SYNC_LOCK = threading.Lock()
_RKN_GIT_SYNC_LAST_TS = 0.0
_USAGE_STATS_LOCK = threading.Lock()
MSK_TZ = ZoneInfo("Europe/Moscow")
MAX_RF_POPULATION = 146_500_000
KNOWN_CALC_TYPES = {"rao_tv", "vois_tv", "rao_radio", "vois_radio", "rao_events", "vois_events"}


def _now_msk() -> datetime:
    return datetime.now(MSK_TZ)


def _registry_title_by_ts(ts: Optional[str] = None) -> str:
    months = {
        1: "январь",
        2: "февраль",
        3: "март",
        4: "апрель",
        5: "май",
        6: "июнь",
        7: "июль",
        8: "август",
        9: "сентябрь",
        10: "октябрь",
        11: "ноябрь",
        12: "декабрь",
    }
    dt: Optional[datetime] = None
    raw = str(ts or "").strip()
    if raw:
        try:
            dt = datetime.fromisoformat(raw.replace("Z", "+00:00"))
        except Exception:
            dt = None
    if dt is None:
        dt = _now_msk()
    try:
        dt = dt.astimezone(MSK_TZ)
    except Exception:
        pass
    return f"Таблица РКН {months.get(int(dt.month), str(dt.month))}"


def _resolve_usage_storage_dir() -> Path:
    raw = str(os.getenv("RKN_USAGE_STORAGE_DIR", "") or "").strip()
    if raw:
        return Path(raw).expanduser()
    data_mount = Path("/data")
    try:
        if data_mount.exists() and data_mount.is_dir():
            return data_mount / "rnk_usage"
    except Exception:
        pass
    app_logs = Path("/app/logs")
    try:
        if app_logs.exists() and app_logs.is_dir():
            return app_logs / "rnk_usage"
    except Exception:
        pass
    return PROJECT_DIR / "calculators" / "rao_tv" / "data"


USAGE_STORAGE_DIR = _resolve_usage_storage_dir()
USAGE_STATS_DB_PATH = USAGE_STORAGE_DIR / "_usage_stats.sqlite"
USAGE_EVENTS_LOG_PATH = USAGE_STORAGE_DIR / "_usage_events.jsonl"
_USAGE_STORAGE_MIGRATED = False
_RKN_UPLOAD_STORAGE_MIGRATED = False


def _cap_radio_population(pop_total: Optional[int], pop_notes: Optional[list[str]]) -> tuple[Optional[int], list[str]]:
    notes = list(pop_notes or [])
    if pop_total is None:
        return None, notes
    try:
        p = int(pop_total)
    except Exception:
        return pop_total, notes
    if p > MAX_RF_POPULATION:
        p_fmt = f"{p:,}".replace(",", " ")
        max_fmt = f"{MAX_RF_POPULATION:,}".replace(",", " ")
        notes.append(
            f"Численность населения по лицензии превышала численность РФ ({p_fmt}); для интерфейса использовано {max_fmt}."
        )
        return MAX_RF_POPULATION, notes
    return p, notes


def _git_run(*args: str, extra_env: Optional[dict[str, str]] = None) -> subprocess.CompletedProcess[str]:
    env = os.environ.copy()
    if extra_env:
        env.update(extra_env)
    return subprocess.run(
        ["git", "-C", str(PROJECT_DIR), *args],
        text=True,
        capture_output=True,
        check=True,
        env=env,
    )


def _git_sync_files(paths: list[Path]) -> str:
    if not RKN_UPLOAD_GIT_COMMIT:
        return "Git-синхронизация отключена."
    if not (PROJECT_DIR / ".git").exists():
        return "Git-репозиторий не найден на сервере."

    rels: list[str] = []
    for p in paths:
        try:
            rels.append(str(p.resolve().relative_to(PROJECT_DIR)))
        except Exception:
            continue
    if not rels:
        return "Обновлены данные на сервере, но файлы не принадлежат репозиторию."

    _git_run("add", "--", *rels)
    diff = subprocess.run(["git", "-C", str(PROJECT_DIR), "diff", "--cached", "--quiet", "--", *rels], text=True)
    if diff.returncode == 0:
        return "Git: изменений для коммита нет."

    ts = _now_msk().strftime("%Y-%m-%d %H:%M:%S")
    msg = f"Update RKN table via web uploader ({ts})"
    env: dict[str, str] = {}
    if RKN_UPLOAD_GIT_AUTHOR_NAME:
        env["GIT_AUTHOR_NAME"] = RKN_UPLOAD_GIT_AUTHOR_NAME
        env["GIT_COMMITTER_NAME"] = RKN_UPLOAD_GIT_AUTHOR_NAME
    if RKN_UPLOAD_GIT_AUTHOR_EMAIL:
        env["GIT_AUTHOR_EMAIL"] = RKN_UPLOAD_GIT_AUTHOR_EMAIL
        env["GIT_COMMITTER_EMAIL"] = RKN_UPLOAD_GIT_AUTHOR_EMAIL

    _git_run("commit", "-m", msg, "--", *rels, extra_env=env)
    if not RKN_UPLOAD_GIT_PUSH:
        return "Git: коммит выполнен, push отключён."

    branch = RKN_UPLOAD_GIT_BRANCH or _git_run("branch", "--show-current").stdout.strip()
    if not branch:
        return "Git: не удалось определить ветку для push."
    _git_run("push", RKN_UPLOAD_GIT_REMOTE, branch)
    return f"Git: изменения отправлены в {RKN_UPLOAD_GIT_REMOTE}/{branch}."


def _git_autosync_if_due(force: bool = False) -> Optional[str]:
    global _RKN_GIT_SYNC_LAST_TS
    if not RKN_UPLOAD_GIT_AUTOSYNC:
        return None
    if not (PROJECT_DIR / ".git").exists():
        return None
    now_ts = time.time()
    if not force and (now_ts - _RKN_GIT_SYNC_LAST_TS) < float(RKN_UPLOAD_GIT_AUTOSYNC_SEC):
        return None
    with _RKN_GIT_SYNC_LOCK:
        now_ts = time.time()
        if not force and (now_ts - _RKN_GIT_SYNC_LAST_TS) < float(RKN_UPLOAD_GIT_AUTOSYNC_SEC):
            return None
        _RKN_GIT_SYNC_LAST_TS = now_ts
        try:
            before = _git_run("rev-parse", "HEAD").stdout.strip()
            branch = RKN_UPLOAD_GIT_BRANCH or _git_run("branch", "--show-current").stdout.strip() or "main"
            _git_run("pull", "--ff-only", RKN_UPLOAD_GIT_REMOTE, branch)
            after = _git_run("rev-parse", "HEAD").stdout.strip()
            if before != after:
                _reset_rkn_caches()
                return f"synced:{before[:7]}->{after[:7]}"
            return "up-to-date"
        except Exception as e:
            return f"sync-error:{e}"


def _reset_rkn_caches() -> None:
    global _RKN_INDEX, _RKN_INDEX_MTIME, _RKN_INDEX_SRC
    _RKN_INDEX = None
    _RKN_INDEX_MTIME = None
    _RKN_INDEX_SRC = None
    _licenses_light_cached.cache_clear()
    refresh_inn_map_if_needed(force=True)
    for fn_name in ("_matching_rkn_rows_cached", "_inn_to_org_map"):
        fn = getattr(rao_mod, fn_name, None)
        if fn is not None and hasattr(fn, "cache_clear"):
            try:
                fn.cache_clear()
            except Exception:
                pass


def _append_rkn_upload_log(entry: dict[str, Any]) -> None:
    try:
        _migrate_rkn_upload_storage_if_needed()
        RKN_UPLOAD_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
        with RKN_UPLOAD_LOG_PATH.open("a", encoding="utf-8") as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception:
        pass


def _rkn_artifacts_status() -> dict[str, Any]:
    items: dict[str, Any] = {}
    for label, path in _active_rkn_artifact_paths().items():
        exists = bool(path and path.exists())
        meta: dict[str, Any] = {"path": str(path) if path else None, "exists": exists}
        if exists:
            st = path.stat()  # type: ignore[union-attr]
            meta["size_bytes"] = int(st.st_size)
            meta["mtime"] = datetime.fromtimestamp(st.st_mtime, tz=MSK_TZ).isoformat(timespec="seconds")
        items[label] = meta
    return items


def _read_rkn_upload_log(limit: int = 20) -> list[dict[str, Any]]:
    if limit < 1:
        limit = 1
    if limit > 200:
        limit = 200
    _migrate_rkn_upload_storage_if_needed()
    if not RKN_UPLOAD_LOG_PATH.exists():
        return []
    rows: list[dict[str, Any]] = []
    try:
        for line in RKN_UPLOAD_LOG_PATH.read_text(encoding="utf-8").splitlines():
            s = line.strip()
            if not s:
                continue
            try:
                obj = json.loads(s)
                if isinstance(obj, dict):
                    rows.append(obj)
            except Exception:
                continue
    except Exception:
        return []
    # Return newest unique records first (protect against duplicated seed/log merges).
    unique_rows: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str]] = set()
    for obj in reversed(rows):
        key = (
            str(obj.get("timestamp") or "").strip(),
            str(obj.get("filename") or "").strip(),
            "1" if bool(obj.get("error")) else "0",
        )
        if key in seen:
            continue
        seen.add(key)
        unique_rows.append(obj)
        if len(unique_rows) >= limit:
            break
    return unique_rows


def _write_rkn_upload_state(entry: dict[str, Any]) -> None:
    try:
        _migrate_rkn_upload_storage_if_needed()
        RKN_UPLOAD_STATE_PATH.parent.mkdir(parents=True, exist_ok=True)
        RKN_UPLOAD_STATE_PATH.write_text(json.dumps(entry, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass


def _read_rkn_upload_state() -> Optional[dict[str, Any]]:
    _migrate_rkn_upload_storage_if_needed()
    if not RKN_UPLOAD_STATE_PATH.exists():
        return None
    try:
        obj = json.loads(RKN_UPLOAD_STATE_PATH.read_text(encoding="utf-8"))
    except Exception:
        return None
    return obj if isinstance(obj, dict) else None


def _set_rkn_upload_job(job_id: str, **fields: Any) -> None:
    with _RKN_UPLOAD_JOBS_LOCK:
        obj = _RKN_UPLOAD_JOBS.get(job_id, {"job_id": job_id})
        obj.update(fields)
        obj["updated_at"] = _now_msk().isoformat(timespec="seconds")
        _RKN_UPLOAD_JOBS[job_id] = obj


def _get_rkn_upload_job(job_id: str) -> Optional[dict[str, Any]]:
    with _RKN_UPLOAD_JOBS_LOCK:
        item = _RKN_UPLOAD_JOBS.get(job_id)
        return dict(item) if item else None


def _init_usage_stats_db() -> None:
    USAGE_STATS_DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    _migrate_usage_storage_if_needed()
    with sqlite3.connect(str(USAGE_STATS_DB_PATH)) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS calc_usage(
                calc_type TEXT PRIMARY KEY,
                count INTEGER NOT NULL DEFAULT 0,
                updated_at TEXT
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS meta(
                key TEXT PRIMARY KEY,
                value TEXT
            )
            """
        )
        conn.commit()


def _migrate_usage_storage_if_needed() -> None:
    global _USAGE_STORAGE_MIGRATED
    if _USAGE_STORAGE_MIGRATED:
        return
    _USAGE_STORAGE_MIGRATED = True
    try:
        if str(USAGE_STATS_DB_PATH.resolve()) != str(LEGACY_USAGE_STATS_DB_PATH.resolve()):
            if not USAGE_STATS_DB_PATH.exists() and LEGACY_USAGE_STATS_DB_PATH.exists():
                USAGE_STATS_DB_PATH.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(LEGACY_USAGE_STATS_DB_PATH, USAGE_STATS_DB_PATH)
    except Exception:
        pass
    try:
        if str(USAGE_EVENTS_LOG_PATH.resolve()) != str(LEGACY_USAGE_EVENTS_LOG_PATH.resolve()):
            if not USAGE_EVENTS_LOG_PATH.exists() and LEGACY_USAGE_EVENTS_LOG_PATH.exists():
                USAGE_EVENTS_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(LEGACY_USAGE_EVENTS_LOG_PATH, USAGE_EVENTS_LOG_PATH)
    except Exception:
        pass


def _migrate_rkn_upload_storage_if_needed() -> None:
    global _RKN_UPLOAD_STORAGE_MIGRATED
    if _RKN_UPLOAD_STORAGE_MIGRATED:
        return
    _RKN_UPLOAD_STORAGE_MIGRATED = True
    try:
        RKN_UPLOAD_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
        if str(RKN_UPLOAD_LOG_PATH.resolve()) != str(LEGACY_RKN_UPLOAD_LOG_PATH.resolve()):
            if not RKN_UPLOAD_LOG_PATH.exists() and LEGACY_RKN_UPLOAD_LOG_PATH.exists():
                shutil.copy2(LEGACY_RKN_UPLOAD_LOG_PATH, RKN_UPLOAD_LOG_PATH)
        if not RKN_UPLOAD_LOG_PATH.exists() and RKN_UPLOAD_LOG_SEED_PATH.exists():
            shutil.copy2(RKN_UPLOAD_LOG_SEED_PATH, RKN_UPLOAD_LOG_PATH)
        elif RKN_UPLOAD_LOG_PATH.exists() and RKN_UPLOAD_LOG_SEED_PATH.exists():
            existing_keys: set[tuple[str, str]] = set()
            try:
                for line in RKN_UPLOAD_LOG_PATH.read_text(encoding="utf-8").splitlines():
                    s = line.strip()
                    if not s:
                        continue
                    obj = json.loads(s)
                    if isinstance(obj, dict):
                        existing_keys.add(
                            (
                                str(obj.get("timestamp") or "").strip(),
                                str(obj.get("filename") or "").strip(),
                            )
                        )
            except Exception:
                existing_keys = set()
            seed_lines: list[str] = []
            try:
                for line in RKN_UPLOAD_LOG_SEED_PATH.read_text(encoding="utf-8").splitlines():
                    s = line.strip()
                    if not s:
                        continue
                    obj = json.loads(s)
                    if not isinstance(obj, dict):
                        continue
                    key = (
                        str(obj.get("timestamp") or "").strip(),
                        str(obj.get("filename") or "").strip(),
                    )
                    if key in existing_keys:
                        continue
                    seed_lines.append(json.dumps(obj, ensure_ascii=False))
            except Exception:
                seed_lines = []
            if seed_lines:
                with RKN_UPLOAD_LOG_PATH.open("a", encoding="utf-8") as f:
                    for line in seed_lines:
                        f.write(line + "\n")
    except Exception:
        pass
    try:
        RKN_UPLOAD_STATE_PATH.parent.mkdir(parents=True, exist_ok=True)
        if str(RKN_UPLOAD_STATE_PATH.resolve()) != str(LEGACY_RKN_UPLOAD_STATE_PATH.resolve()):
            if not RKN_UPLOAD_STATE_PATH.exists() and LEGACY_RKN_UPLOAD_STATE_PATH.exists():
                shutil.copy2(LEGACY_RKN_UPLOAD_STATE_PATH, RKN_UPLOAD_STATE_PATH)
    except Exception:
        pass


def _increment_usage(calc_type: str) -> None:
    ct = str(calc_type or "").strip().lower() or "unknown"
    ts = _now_msk().isoformat(timespec="seconds")
    with _USAGE_STATS_LOCK:
        _init_usage_stats_db()
        with sqlite3.connect(str(USAGE_STATS_DB_PATH)) as conn:
            conn.execute(
                """
                INSERT INTO calc_usage(calc_type, count, updated_at)
                VALUES(?, 1, ?)
                ON CONFLICT(calc_type)
                DO UPDATE SET count = count + 1, updated_at = excluded.updated_at
                """,
                (ct, ts),
            )
            conn.execute(
                """
                INSERT INTO meta(key, value) VALUES('updated_at', ?)
                ON CONFLICT(key) DO UPDATE SET value = excluded.value
                """,
                (ts,),
            )
            conn.commit()
    try:
        USAGE_EVENTS_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
        with USAGE_EVENTS_LOG_PATH.open("a", encoding="utf-8") as f:
            f.write(json.dumps({"timestamp": ts, "calc_type": ct, "source": "runtime_increment"}, ensure_ascii=False) + "\n")
    except Exception:
        pass


def _read_usage_stats() -> dict[str, Any]:
    with _USAGE_STATS_LOCK:
        _init_usage_stats_db()
        with sqlite3.connect(str(USAGE_STATS_DB_PATH)) as conn:
            rows = conn.execute(
                "SELECT calc_type, count, updated_at FROM calc_usage ORDER BY count DESC, calc_type ASC"
            ).fetchall()
            total = int(sum(int(r[1] or 0) for r in rows))
            updated_at_row = conn.execute("SELECT value FROM meta WHERE key='updated_at'").fetchone()
    items = [
        {
            "calc_type": str(r[0] or ""),
            "count": int(r[1] or 0),
            "updated_at": str(r[2] or ""),
        }
        for r in rows
    ]
    return {
        "total": total,
        "items": items,
        "updated_at": str(updated_at_row[0]) if updated_at_row and updated_at_row[0] else None,
        "db_path": str(USAGE_STATS_DB_PATH),
        "storage_dir": str(USAGE_STORAGE_DIR),
    }


def _usage_backfill_patterns() -> list[str]:
    raw = str(USAGE_BACKFILL_LOG_GLOBS or "").strip()
    if not raw:
        return []
    return [p.strip() for p in re.split(r"[;\n,]+", raw) if p.strip()]


def _usage_backfill_resolve_files(patterns: list[str]) -> list[Path]:
    files: list[Path] = []
    seen: set[str] = set()
    for pattern in patterns:
        for raw_path in glob.glob(pattern):
            p = Path(raw_path).expanduser()
            try:
                rp = str(p.resolve())
            except Exception:
                rp = str(p)
            if rp in seen:
                continue
            seen.add(rp)
            if p.exists() and p.is_file():
                files.append(Path(rp))
    files.sort(key=lambda x: str(x))
    return files


def _usage_backfill_iter_lines(path: Path):
    opener = gzip.open if path.suffix == ".gz" else open
    try:
        with opener(path, "rt", encoding="utf-8", errors="ignore") as fh:
            for line in fh:
                yield line.rstrip("\n")
    except Exception:
        return


def _usage_normalize_calc_type(raw: Any) -> Optional[str]:
    ct = str(raw or "").strip().lower()
    if ct == "vois_rv":
        ct = "vois_radio"
    return ct if ct in KNOWN_CALC_TYPES else None


def _usage_backfill_detect_calc_type(line: str) -> Optional[str]:
    s = str(line or "")
    if not s:
        return None
    sl = s.lower()
    # 1) явные значения calc_type в query/json/log payload
    m = re.search(r"(?:\"|')?calc_type(?:\"|')?\s*[:=]\s*(?:\"|')([a-z_]+)(?:\"|')", sl)
    if m:
        ct = _usage_normalize_calc_type(m.group(1))
        if ct:
            return ct

    # 2) fallback по url, если в логах есть route с calc_type в query
    if "/api/calc" in sl and "calc_type=" in sl:
        mq = re.search(r"calc_type=([a-z_]+)", sl)
        if mq:
            ct = _usage_normalize_calc_type(mq.group(1))
            if ct in {"rao_tv", "vois_tv", "rao_radio", "vois_radio"}:
                return ct
    if "/api/events/calc" in sl and "calc_type=" in sl:
        mq = re.search(r"calc_type=([a-z_]+)", sl)
        if mq:
            ct = _usage_normalize_calc_type(mq.group(1))
            if ct in {"rao_events", "vois_events"}:
                return ct
    return None


def _usage_backfill_detect_heuristic(line: str) -> Optional[str]:
    s = str(line or "")
    if not s:
        return None
    sl = s.lower()
    # Эвристика по запросам списка лицензий: в них calc_type всегда в query.
    # Используем только как приблизительный источник, когда явных calc_type нет.
    if "/api/licenses" in sl and "calc_type=" in sl:
        mq = re.search(r"calc_type=([a-z_]+)", sl)
        if mq:
            ct = _usage_normalize_calc_type(mq.group(1))
            if ct in {"rao_tv", "vois_tv", "rao_radio", "vois_radio"}:
                return ct
    return None


def _usage_apply_floor_counts(counts: dict[str, int]) -> None:
    if not counts:
        return
    ts = _now_msk().isoformat(timespec="seconds")
    with _USAGE_STATS_LOCK:
        _init_usage_stats_db()
        with sqlite3.connect(str(USAGE_STATS_DB_PATH)) as conn:
            for ct, total_count in counts.items():
                norm_ct = _usage_normalize_calc_type(ct)
                if not norm_ct:
                    continue
                n = int(total_count or 0)
                if n <= 0:
                    continue
                conn.execute(
                    """
                    INSERT INTO calc_usage(calc_type, count, updated_at)
                    VALUES(?, ?, ?)
                    ON CONFLICT(calc_type)
                    DO UPDATE SET
                      count = CASE WHEN calc_usage.count < excluded.count THEN excluded.count ELSE calc_usage.count END,
                      updated_at = excluded.updated_at
                    """,
                    (norm_ct, n, ts),
                )
            conn.execute(
                """
                INSERT INTO meta(key, value) VALUES('updated_at', ?)
                ON CONFLICT(key) DO UPDATE SET value = excluded.value
                """,
                (ts,),
            )
            conn.commit()


def _usage_backfill_from_logs(patterns: Optional[list[str]] = None, reset: bool = False, dry_run: bool = False) -> dict[str, Any]:
    use_patterns = patterns if patterns else _usage_backfill_patterns()
    files = _usage_backfill_resolve_files(use_patterns)
    delta_explicit: dict[str, int] = {}
    delta_heuristic: dict[str, int] = {}
    scanned_lines = 0
    matched_lines = 0
    matched_explicit = 0
    matched_heuristic = 0
    for p in files:
        for line in _usage_backfill_iter_lines(p):
            scanned_lines += 1
            ct = _usage_backfill_detect_calc_type(line)
            if ct:
                matched_lines += 1
                matched_explicit += 1
                delta_explicit[ct] = int(delta_explicit.get(ct, 0)) + 1
                continue
            ht = _usage_backfill_detect_heuristic(line)
            if ht:
                matched_lines += 1
                matched_heuristic += 1
                delta_heuristic[ht] = int(delta_heuristic.get(ht, 0)) + 1
    # Источники комплементарны: явные записи обычно точнее, эвристика
    # закрывает кейсы, где calc_type не логируется в теле /api/calc.
    # Берём максимум по каждому калькулятору.
    delta: dict[str, int] = {}
    for ct, n in delta_explicit.items():
        if n > int(delta.get(ct, 0)):
            delta[ct] = int(n)
    for ct, n in delta_heuristic.items():
        if n > int(delta.get(ct, 0)):
            delta[ct] = int(n)
    # Никогда не обнуляем статистику: reset игнорируем принудительно.
    # Данные из логов трактуем как "минимально гарантированный общий счётчик".
    if not dry_run and delta:
        _usage_apply_floor_counts(delta)
    return {
        "patterns": use_patterns,
        "files_scanned": [str(p) for p in files],
        "files_count": len(files),
        "lines_scanned": scanned_lines,
        "lines_matched": matched_lines,
        "matched_explicit": matched_explicit,
        "matched_heuristic": matched_heuristic,
        "heuristic_used": bool(delta_heuristic),
        "delta_explicit": delta_explicit,
        "delta_heuristic": delta_heuristic,
        "delta": delta,
        "reset_applied": False,
        "dry_run": bool(dry_run),
}


def _usage_event_log_files() -> list[Path]:
    files: list[Path] = []
    seen: set[str] = set()
    base_dirs: list[Path] = [USAGE_EVENTS_LOG_PATH.parent]
    try:
        legacy_dir = LEGACY_USAGE_EVENTS_LOG_PATH.parent.resolve()
        if legacy_dir not in base_dirs:
            base_dirs.append(legacy_dir)
    except Exception:
        pass
    patterns: list[str] = []
    for bdir in base_dirs:
        patterns.append(str((bdir / "_usage_events*.jsonl").resolve()))
        patterns.append(str((bdir / "_usage_events*.jsonl.*").resolve()))
    for pattern in patterns:
        for raw_path in glob.glob(pattern):
            p = Path(raw_path).expanduser()
            if not p.exists() or not p.is_file():
                continue
            try:
                rp = str(p.resolve())
            except Exception:
                rp = str(p)
            if rp in seen:
                continue
            seen.add(rp)
            files.append(Path(rp))
    files.sort(key=lambda x: str(x))
    return files


def _usage_counts_from_event_logs(files: list[Path]) -> dict[str, int]:
    out: dict[str, int] = {}
    for p in files:
        for line in _usage_backfill_iter_lines(p):
            s = str(line or "").strip()
            if not s:
                continue
            try:
                obj = json.loads(s)
            except Exception:
                continue
            if not isinstance(obj, dict):
                continue
            ct = _usage_normalize_calc_type(obj.get("calc_type"))
            if not ct:
                continue
            out[ct] = int(out.get(ct, 0)) + 1
    return out


def _usage_reconcile_from_logs_monotonic(force: bool = False) -> Optional[dict[str, Any]]:
    # Источник №1: внутренний event-log расчётов (наиболее точный).
    event_files = _usage_event_log_files()
    event_counts = _usage_counts_from_event_logs(event_files)

    # Источник №2: внешние логи API (fallback/добавка).
    patterns = _usage_backfill_patterns()
    generic_files = _usage_backfill_resolve_files(patterns)
    explicit_counts: dict[str, int] = {}
    heuristic_counts: dict[str, int] = {}
    for p in generic_files:
        for line in _usage_backfill_iter_lines(p):
            ct = _usage_backfill_detect_calc_type(line)
            if ct:
                explicit_counts[ct] = int(explicit_counts.get(ct, 0)) + 1
                continue
            ht = _usage_backfill_detect_heuristic(line)
            if ht:
                heuristic_counts[ht] = int(heuristic_counts.get(ht, 0)) + 1
    generic_counts: dict[str, int] = {}
    for ct, n in explicit_counts.items():
        if n > int(generic_counts.get(ct, 0)):
            generic_counts[ct] = int(n)
    for ct, n in heuristic_counts.items():
        if n > int(generic_counts.get(ct, 0)):
            generic_counts[ct] = int(n)

    merged: dict[str, int] = {}
    for src in (event_counts, generic_counts):
        for ct, n in src.items():
            if n > int(merged.get(ct, 0)):
                merged[ct] = int(n)

    # Защита от лишней работы: если сигнатура логов не изменилась, второй раз не сканируем/не пишем.
    def _sig_part(path: Path) -> str:
        try:
            st = path.stat()
            return f"{path}:{int(st.st_size)}:{int(st.st_mtime_ns)}"
        except Exception:
            return f"{path}:0:0"

    all_files = sorted({str(p) for p in event_files + generic_files})
    sig = "|".join([_sig_part(Path(p)) for p in all_files])

    with _USAGE_STATS_LOCK:
        _init_usage_stats_db()
        with sqlite3.connect(str(USAGE_STATS_DB_PATH)) as conn:
            old_sig = _meta_get(conn, "usage_reconcile_sig_v1") or ""
            if (not force) and old_sig == sig:
                return None
            _meta_set(conn, "usage_reconcile_sig_v1", sig)
            _meta_set(conn, "usage_reconcile_ts_v1", _now_msk().isoformat(timespec="seconds"))
            conn.commit()

    if merged:
        _usage_apply_floor_counts(merged)

    return {
        "event_files": [str(p) for p in event_files],
        "event_counts": event_counts,
        "generic_files": [str(p) for p in generic_files],
        "generic_counts": generic_counts,
        "merged_counts": merged,
    }


def _usage_scan_diagnostics() -> dict[str, Any]:
    patterns = _usage_backfill_patterns()
    generic_files = _usage_backfill_resolve_files(patterns)
    event_files = _usage_event_log_files()
    return {
        "usage_storage_dir": str(USAGE_STORAGE_DIR),
        "usage_db_path": str(USAGE_STATS_DB_PATH),
        "usage_event_log_path": str(USAGE_EVENTS_LOG_PATH),
        "legacy_db_path": str(LEGACY_USAGE_STATS_DB_PATH),
        "legacy_event_log_path": str(LEGACY_USAGE_EVENTS_LOG_PATH),
        "patterns": patterns,
        "generic_files_count": len(generic_files),
        "event_files_count": len(event_files),
        "generic_files": [str(p) for p in generic_files[:20]],
        "event_files": [str(p) for p in event_files[:20]],
    }


def _meta_get(conn: sqlite3.Connection, key: str) -> Optional[str]:
    row = conn.execute("SELECT value FROM meta WHERE key=?", (str(key),)).fetchone()
    if not row:
        return None
    val = row[0]
    return str(val) if val is not None else None


def _meta_set(conn: sqlite3.Connection, key: str, value: str) -> None:
    conn.execute(
        """
        INSERT INTO meta(key, value) VALUES(?, ?)
        ON CONFLICT(key) DO UPDATE SET value = excluded.value
        """,
        (str(key), str(value)),
    )


def _usage_backfill_auto_once() -> Optional[dict[str, Any]]:
    with _USAGE_STATS_LOCK:
        _init_usage_stats_db()
        with sqlite3.connect(str(USAGE_STATS_DB_PATH)) as conn:
            done = _meta_get(conn, "backfill_once_done_v2")
            if done:
                return None
            _meta_set(conn, "backfill_once_done_v2", _now_msk().isoformat(timespec="seconds"))
            conn.commit()
    try:
        report = _usage_backfill_from_logs(patterns=None, dry_run=False)
        return report
    except Exception:
        return None


def _load_admin_users() -> dict[str, dict[str, str]]:
    default_users = [
        {"login": "ivan.kunitsyn", "name": "Иван Куницын", "password": "<REDACTED>"},
        {"login": "igor.kunitsyn", "name": "Игорь Куницын", "password": "<REDACTED>"},
    ]
    payload: Any = None
    if RKN_ADMIN_USERS_JSON:
        try:
            payload = json.loads(RKN_ADMIN_USERS_JSON)
        except Exception:
            payload = None
    if payload is None:
        payload = default_users

    users: dict[str, dict[str, str]] = {}
    if isinstance(payload, dict):
        for login, password in payload.items():
            lg = str(login or "").strip()
            pw = str(password or "").strip()
            if not lg or not pw:
                continue
            users[lg.lower()] = {"login": lg, "name": lg, "password": pw}
    elif isinstance(payload, list):
        for item in payload:
            if not isinstance(item, dict):
                continue
            lg = str(item.get("login") or "").strip()
            pw = str(item.get("password") or "").strip()
            nm = str(item.get("name") or lg).strip()
            if not lg or not pw:
                continue
            users[lg.lower()] = {"login": lg, "name": nm or lg, "password": pw}
    return users


def _check_admin_credentials(login: str, password: str) -> tuple[bool, Optional[dict[str, str]]]:
    lg = str(login or "").strip()
    pw = str(password or "").strip()
    if not lg or not pw:
        return False, None
    users = _load_admin_users()
    user = users.get(lg.lower())
    if not user:
        return False, None
    if str(user.get("password") or "") != pw:
        return False, None
    return True, user


def _run_rkn_upload_job(
    job_id: str,
    raw_tmp: Path,
    original_filename: str,
    uploader: Optional[str],
    ip: str,
    user_agent: str,
) -> None:
    _set_rkn_upload_job(job_id, status="processing", message="Обработка таблицы…")
    with _RKN_UPLOAD_LOCK:
        backup_dir = RKN_UPLOAD_STORAGE_DIR / "_backup"
        backup_dir.mkdir(parents=True, exist_ok=True)
        ts = _now_msk().strftime("%Y%m%d_%H%M%S")

        output_paths = _rkn_upload_output_paths()
        output_xlsx = output_paths["xlsx"]
        output_sqlite = output_paths["sqlite"]
        output_inn_csv = output_paths["inn_csv"]
        output_xlsx.parent.mkdir(parents=True, exist_ok=True)
        output_sqlite.parent.mkdir(parents=True, exist_ok=True)
        output_inn_csv.parent.mkdir(parents=True, exist_ok=True)
        backup_xlsx = backup_dir / f"{ts}_{output_xlsx.name}"
        backup_sqlite = backup_dir / f"{ts}_{output_sqlite.name}"
        backup_inn_csv = backup_dir / f"{ts}_inn_name.csv"
        try:
            content_size = int(raw_tmp.stat().st_size) if raw_tmp.exists() else 0
            if output_xlsx.exists():
                shutil.copy2(output_xlsx, backup_xlsx)
            if output_sqlite.exists():
                shutil.copy2(output_sqlite, backup_sqlite)
            if output_inn_csv.exists():
                shutil.copy2(output_inn_csv, backup_inn_csv)

            requested_population_mode = RKN_UPLOAD_POPULATION_MODE
            used_population_mode = requested_population_mode
            mode_note = ""
            try:
                stats = clean_rkn_xlsx(
                    src_path=raw_tmp,
                    dst_path=output_xlsx,
                    allow_empty_status=RKN_UPLOAD_ALLOW_EMPTY_STATUS,
                    population_mode=requested_population_mode,
                )
            except Exception as mode_error:
                if requested_population_mode == POPULATION_MODE_YANDEX:
                    used_population_mode = "legacy"
                    mode_note = f"Yandex-нормализация недоступна, применён fallback={used_population_mode}: {mode_error}"
                    stats = clean_rkn_xlsx(
                        src_path=raw_tmp,
                        dst_path=output_xlsx,
                        allow_empty_status=RKN_UPLOAD_ALLOW_EMPTY_STATUS,
                        population_mode=used_population_mode,
                    )
                else:
                    raise

            sqlite_rows = build_rkn_sqlite(output_xlsx, output_sqlite)
            inn_rows = build_inn_name_csv(output_xlsx, output_inn_csv)
            verify = verify_rkn_artifacts(
                output_xlsx,
                sqlite_path=output_sqlite,
                inn_csv_path=output_inn_csv,
            )
            _reset_rkn_caches()
            git_report = _git_sync_files([output_xlsx, output_sqlite, output_inn_csv])

            log_entry = {
                "timestamp": _now_msk().isoformat(timespec="seconds"),
                "filename": _registry_title_by_ts(),
                "source_filename": original_filename or "uploaded.xlsx",
                "bytes": content_size,
                "uploader": str(uploader or "").strip() or None,
                "ip": ip,
                "user_agent": user_agent[:300],
                "population_mode": RKN_UPLOAD_POPULATION_MODE,
                "population_mode_used": used_population_mode,
                "population_mode_note": mode_note,
                "stats": stats.as_text(),
                "verify": verify.as_text(),
                "sqlite_rows": sqlite_rows,
                "inn_rows": inn_rows,
                "artifacts": {
                    "xlsx": str(output_xlsx),
                    "sqlite": str(output_sqlite),
                    "inn_csv": str(output_inn_csv),
                },
                "git": git_report,
            }
            _append_rkn_upload_log(log_entry)
            _write_rkn_upload_state(log_entry)
            _set_rkn_upload_job(
                job_id,
                status="done",
                ok=True,
                message="Таблица РКН обновлена.",
                data={
                    "ok": True,
                    "message": "Таблица РКН обновлена.",
                    "population_mode": RKN_UPLOAD_POPULATION_MODE,
                    "population_mode_used": used_population_mode,
                    "population_mode_note": mode_note,
                    "stats": stats.as_text(),
                    "verify": verify.as_text(),
                    "sqlite_rows": sqlite_rows,
                    "inn_rows": inn_rows,
                    "git": git_report,
                    "backup_dir": str(backup_dir),
                    "log_entry": log_entry,
                    "log_path": str(RKN_UPLOAD_LOG_PATH),
                    "artifacts": _rkn_artifacts_status(),
                },
            )
        except Exception as e:
            err = f"Ошибка обновления таблицы: {e}"
            _append_rkn_upload_log(
                {
                    "timestamp": _now_msk().isoformat(timespec="seconds"),
                    "filename": _registry_title_by_ts(),
                    "source_filename": original_filename or "uploaded.xlsx",
                    "uploader": str(uploader or "").strip() or None,
                    "ip": ip,
                    "error": f"{e}",
                }
            )
            _set_rkn_upload_job(job_id, status="error", ok=False, error=err)
        finally:
            try:
                if raw_tmp.exists():
                    raw_tmp.unlink()
            except Exception:
                pass


def _is_active_status(raw_status: Any) -> bool:
    s = str(raw_status or "").strip().lower()
    if not s:
        return True
    negative_markers = (
        "недейств",
        "аннули",
        "прекращ",
        "приостанов",
        "истек",
    )
    if any(x in s for x in negative_markers):
        return False
    return True


def _sorted_unique(values: Any) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for v in (values or []):
        s = str(v or "").strip()
        if not s or s in seen:
            continue
        seen.add(s)
        out.append(s)
    out.sort(key=lambda x: (len(x), x))
    return out


def _format_smi_display(values: Any) -> str:
    raw = _sorted_unique(values)
    if not raw:
        return ""
    by_base: dict[str, str] = {}
    for item in raw:
        base = item.split("—", 1)[0].strip()
        prev = by_base.get(base)
        # Предпочитаем вариант с часами вещания.
        if prev is None or ("—" in item and "—" not in prev):
            by_base[base] = item
    merged = sorted(by_base.values(), key=lambda x: (len(x), x))
    return " / ".join(merged)


def _collect_smi_from_raw(lic: dict, smi14: str, smi: str) -> None:
    parsed_channels = rao_mod.split_channel_tokens(smi14) or rao_mod.split_channel_tokens(smi)
    if parsed_channels:
        for ch_name, ch_hours in parsed_channels:
            lic["smi_values"].add(ch_name)
            if ch_hours is not None:
                lic["smi_display_values"].add(f"{ch_name} — {str(ch_hours).replace('.', ',')} ч")
            else:
                lic["smi_display_values"].add(ch_name)
            lic["smi_radio_flags"].add(bool(rao_mod.is_radio_channel_name(ch_name)))
        return
    if smi14:
        lic["smi_values"].add(smi14)
        lic["smi_display_values"].add(smi14)
        lic["smi_radio_flags"].add(bool(rao_mod.is_radio_channel_name(smi14)))
        return
    if smi:
        lic["smi_values"].add(smi)
        lic["smi_display_values"].add(smi)
        lic["smi_radio_flags"].add(bool(rao_mod.is_radio_channel_name(smi)))


def _activity_is_radio(raw: Any) -> bool:
    s = str(raw or "").strip().lower()
    if not s:
        return False
    return any(x in s for x in ("радиовещ", "радиоканал", "радиопрограм"))


def _activity_is_tv(raw: Any) -> bool:
    s = str(raw or "").strip().lower()
    if not s:
        return False
    return any(x in s for x in ("телевещ", "телеканал", "телепрограм", "телепередач"))


def _load_rao_module():
    from . import rao as mod  # type: ignore
    return mod


rao_mod = _load_rao_module()
run_calc_capture = rao_mod.run_calc_capture
parse_inn = rao_mod.parse_inn
get_org_name_by_inn = rao_mod.get_org_name_by_inn
fix_mojibake = rao_mod.fix_mojibake
parse_population_override = rao_mod.parse_population_override

RAO_DIR = Path(rao_mod.__file__).resolve().parent

def find_rkn_xlsx() -> Path:
    env_path = os.getenv("RKN_XLSX_PATH", "").strip()
    if env_path:
        p = Path(env_path)
        if p.exists():
            return p
    candidates = [
        TV_DATA_DIR / "Таблица РКН slim.xlsx",
        TV_DATA_DIR / "Таблица РКН очищенная.xlsx",
        TV_DATA_DIR / "Таблица РКН.xlsx",
        TV_DATA_DIR / "Таблица РКН (2).xlsx",
        BASE_DIR / "Таблица РКН slim.xlsx",
        BASE_DIR / "Таблица РКН очищенная.xlsx",
        BASE_DIR / "Таблица РКН.xlsx",
        BASE_DIR / "Таблица РКН (2).xlsx",
        RAO_DIR / "Таблица РКН slim.xlsx",
        RAO_DIR / "Таблица РКН очищенная.xlsx",
        RAO_DIR / "Таблица РКН.xlsx",
        RAO_DIR / "Таблица РКН (2).xlsx",
        RKN_BOT_DIR / "Таблица РКН slim.xlsx",
        RKN_BOT_DIR / "Таблица РКН очищенная.xlsx",
        RKN_BOT_DIR / "Таблица РКН.xlsx",
        Path.cwd() / "Таблица РКН slim.xlsx",
        Path.cwd() / "Таблица РКН очищенная.xlsx",
        Path.cwd() / "Таблица РКН.xlsx",
        Path.cwd() / "Таблица РКН (2).xlsx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(
        "Не найден файл 'Таблица РКН.xlsx'. "
        f"Пробовал: {[str(c) for c in candidates]}"
    )

def find_rkn_db() -> Optional[Path]:
    env_path = os.getenv("RKN_DB_PATH", "").strip()
    if env_path:
        p = Path(env_path)
        if p.exists():
            return p
        # Если путь к БД задан явно, не переключаемся на bundled fallback.
        # Иначе можно незаметно использовать устаревшую локальную БД из образа.
        return None
    candidates = [
        TV_DATA_DIR / "Таблица РКН.sqlite",
        TV_DATA_DIR / "Таблица РКН.db",
        BASE_DIR / "Таблица РКН.sqlite",
        BASE_DIR / "Таблица РКН.db",
        RAO_DIR / "Таблица РКН.sqlite",
        RAO_DIR / "Таблица РКН.db",
        RKN_BOT_DIR / "Таблица РКН.sqlite",
        RKN_BOT_DIR / "Таблица РКН.db",
        Path.cwd() / "Таблица РКН.sqlite",
        Path.cwd() / "Таблица РКН.db",
    ]
    for p in candidates:
        if p.exists():
            return p
    return None


def _active_rkn_artifact_paths() -> dict[str, Optional[Path]]:
    xlsx: Optional[Path] = None
    sqlite: Optional[Path] = None
    inn_csv: Optional[Path] = None
    try:
        xlsx = find_rkn_xlsx()
    except Exception:
        xlsx = None
    try:
        sqlite = find_rkn_db()
    except Exception:
        sqlite = None
    try:
        inn_csv = _resolve_inn_csv_path()
    except Exception:
        inn_csv = None
    return {"xlsx": xlsx, "sqlite": sqlite, "inn_csv": inn_csv}

def find_vars_xlsx(calc_type: str = "rao_tv") -> Path:
    calc_type = (calc_type or "rao_tv").strip().lower()
    if calc_type == "vois_tv":
        env_path = os.getenv("RKN_VOIS_VARS_XLSX_PATH", "").strip()
        if env_path:
            p = Path(env_path)
            if p.exists():
                return p
        vois_candidates = [
            VOIS_TV_DATA_DIR / "Переменные из ставок ВОИС.xlsx",
            VOIS_TV_DATA_DIR / "Переменные из ставок.xlsx",
            BASE_DIR / "Переменные из ставок ВОИС.xlsx",
            PROJECT_DIR / "Переменные из ставок ВОИС.xlsx",
        ]
        for p in vois_candidates:
            if p.exists():
                return p

    env_path = os.getenv("RKN_VARS_XLSX_PATH", "").strip()
    if env_path:
        p = Path(env_path)
        if p.exists():
            return p
    candidates = [
        TV_DATA_DIR / "Переменные из ставок.xlsx",
        BASE_DIR / "Переменные из ставок.xlsx",
        RAO_DIR / "Переменные из ставок.xlsx",
        Path.cwd() / "Переменные из ставок.xlsx",
    ]
    for p in candidates:
        if p.exists():
            return p
    name = "Переменные из ставок ВОИС.xlsx" if calc_type == "vois_tv" else "Переменные из ставок.xlsx"
    raise FileNotFoundError(
        f"Не найден файл '{name}'. "
        f"Пробовал: {[str(c) for c in candidates]}"
    )


def _iter_rkn_rows_light(rkn_xlsx: Path):
    wb = openpyxl.load_workbook(rkn_xlsx, read_only=True, data_only=True)
    ws = wb.active
    header_raw = list(next(ws.iter_rows(min_row=1, max_row=1, values_only=True)))
    while header_raw and (header_raw[-1] is None or str(header_raw[-1]).strip() == ""):
        header_raw.pop()
    header = header_raw
    max_col = len(header)
    it = ws.iter_rows(min_row=2, max_col=max_col, values_only=True)
    return header, it


def _has_freq_value(raw: Any) -> bool:
    return bool(str(raw or "").strip())


def _normalize_radio_population_rows(
    data: dict,
    *,
    runtime: bool,
) -> tuple[Optional[int], list[str]]:
    pop_rows = list(data.get("pop_rows") or [])
    if not pop_rows:
        return None, []

    rows_with_freq = [row for row in pop_rows if _has_freq_value(row.get("freq"))]
    if not rows_with_freq:
        return None, []

    total = 0
    notes: list[str] = []
    for row in rows_with_freq:
        raw = row.get("population")
        if runtime:
            pop_int, row_notes = rao_mod.parse_population_runtime(
                raw,
                media_raw=row.get("brcst_descr") or data.get("media_raw"),
                region_name=row.get("region_name_full"),
                region_text=row.get("region_text"),
            )
        else:
            pop_int, row_notes = rao_mod.parse_population(raw)
        if pop_int is not None:
            total += int(pop_int)
        if row_notes:
            notes.extend(row_notes)
    if total <= 0:
        return None, notes[:2]
    return int(total), notes[:2]


def _normalize_license_population_runtime(data: dict) -> tuple[Optional[int], list[str]]:
    flags = data.get("smi_radio_flags") or set()
    media_l = str(data.get("media_raw") or "").lower()
    if _activity_is_radio(data.get("licensed_activity")) or ("радио" in media_l) or ("радиоканал" in media_l) or ("радиовещ" in media_l) or (flags and all(bool(x) for x in flags)):
        radio_total, radio_notes = _normalize_radio_population_rows(data, runtime=True)
        if radio_total is not None:
            return radio_total, radio_notes
    raw_values = list(data.get("pop_raw_values") or [])
    if not raw_values:
        pop_values = sorted(data.get("pop_values") or [])
        if not pop_values:
            return None, []
        return int(max(set(pop_values))), list(data.get("pop_notes") or [])[:2]

    pop_values: set[int] = set()
    pop_notes: list[str] = []
    for raw in raw_values:
        pop_int, notes = rao_mod.parse_population_runtime(
            raw,
            media_raw=data.get("media_raw"),
            region_name=data.get("region_name_full"),
            region_text=data.get("region_text"),
        )
        if pop_int is not None:
            pop_values.add(int(pop_int))
        if notes:
            pop_notes.extend(notes)
    if not pop_values:
        return None, pop_notes[:2]
    if len(pop_values) > 1:
        pop_notes.append("Обнаружено несколько значений населения по лицензии; для расчёта использовано максимальное значение.")
    return int(max(pop_values)), pop_notes[:2]


def _normalize_license_population_fast(data: dict) -> tuple[Optional[int], list[str]]:
    """
    Быстрая нормализация населения без сетевых вызовов.
    Используется в /api/licenses, чтобы не ловить 504 на шаге выбора лицензий.
    """
    flags = data.get("smi_radio_flags") or set()
    media_l = str(data.get("media_raw") or "").lower()
    if _activity_is_radio(data.get("licensed_activity")) or ("радио" in media_l) or ("радиоканал" in media_l) or ("радиовещ" in media_l) or (flags and all(bool(x) for x in flags)):
        radio_total, radio_notes = _normalize_radio_population_rows(data, runtime=False)
        if radio_total is not None:
            return radio_total, radio_notes
    raw_values = list(data.get("pop_raw_values") or [])
    if not raw_values:
        pop_values = sorted(data.get("pop_values") or [])
        if not pop_values:
            return None, []
        return int(max(set(pop_values))), list(data.get("pop_notes") or [])[:2]

    pop_values: set[int] = set()
    pop_notes: list[str] = []
    for raw in raw_values:
        pop_int, notes = rao_mod.parse_population(raw)
        if pop_int is not None:
            pop_values.add(int(pop_int))
        if notes:
            pop_notes.extend(notes)
    if not pop_values:
        return None, pop_notes[:2]
    if len(pop_values) > 1:
        pop_notes.append("Обнаружено несколько значений населения по лицензии; для расчёта использовано максимальное значение.")
    return int(max(pop_values)), pop_notes[:2]


@lru_cache(maxsize=64)
def _licenses_light_cached(rkn_path: str, mtime: float, inn: str, include_radio: bool):
    rkn_xlsx = Path(rkn_path)
    header, it = _iter_rkn_rows_light(rkn_xlsx)
    idx = {h: i for i, h in enumerate(header)}

    def get(row, col):
        j = idx.get(col)
        if j is None:
            return None
        if j >= len(row):
            return None
        return row[j]

    by_license = {}
    for row in it:
        row_inn = str(get(row, "ns1:inn") or "").strip()
        if row_inn != inn:
            continue

        status = get(row, "ns1:status")
        if not _is_active_status(status):
            continue

        lic_id = str(get(row, "ns1:license_num") or "").strip()
        if not lic_id:
            continue

        org_name = str(get(row, "ns1:org_name") or "").strip()
        sreda = str(get(row, "ns1:sreda") or "").strip()
        pop_raw = get(row, "ns1:population")
        region_name = str(get(row, "ns1:region_name_full") or "").strip()
        region_text = str(get(row, "ns1:region_text") or "").strip()
        smi14 = str(get(row, "ns1:smi_name14") or "").strip()
        smi = str(get(row, "ns1:smi_name") or "").strip()
        licensed_activity = str(get(row, "ns1:licensed_activity") or "").strip()
        license_date = str(get(row, "ns1:license_date") or "").strip()
        service_start_date = str(get(row, "ns1:service_start_date") or "").strip()
        freq = get(row, "ns1:freq")
        brcst_descr = str(get(row, "ns1:brcst_descr") or "").strip()

        lic = by_license.setdefault(lic_id, {
            "org_name": org_name,
            "media_raw": sreda,
            "licensed_activity": licensed_activity,
            "license_date": license_date,
            "service_start_date": service_start_date,
            "region_name_full": region_name,
            "region_text": region_text,
            "pop_values": set(),
            "pop_raw_values": set(),
            "pop_notes": [],
            "pop_rows": [],
            "smi_values": set(),
            "smi_display_values": set(),
            "smi_radio_flags": set(),
        })

        if pop_raw is not None and str(pop_raw).strip() != "":
            lic["pop_raw_values"].add(str(pop_raw).strip())
            lic["pop_rows"].append({
                "population": pop_raw,
                "freq": freq,
                "brcst_descr": brcst_descr or sreda,
                "region_name_full": region_name,
                "region_text": region_text,
            })

        if not lic["media_raw"] and sreda:
            lic["media_raw"] = sreda
        if not lic.get("licensed_activity") and licensed_activity:
            lic["licensed_activity"] = licensed_activity
        if not lic.get("license_date") and license_date:
            lic["license_date"] = license_date
        if not lic.get("service_start_date") and service_start_date:
            lic["service_start_date"] = service_start_date
        _collect_smi_from_raw(lic, smi14, smi)

    items = []
    for lic_id, data in by_license.items():
        flags = data.get("smi_radio_flags") or set()
        licensed_activity = str(data.get("licensed_activity") or "")
        if (not include_radio) and (_activity_is_radio(licensed_activity) or (flags and all(flags))):
            continue
        pop_total, pop_notes = _normalize_license_population_fast(data)
        media_raw = data.get("media_raw") or ""
        media_class = rao_mod.normalize_media(media_raw)
        smi_names = _sorted_unique(data.get("smi_values"))
        smi_display_list = _sorted_unique(data.get("smi_display_values"))
        smi_name = smi_names[0] if smi_names else ""
        smi_display = _format_smi_display(smi_display_list) if smi_display_list else smi_name
        items.append({
            "license_id": lic_id,
            "media_raw": media_raw,
            "licensed_activity": licensed_activity,
            "license_date": str(data.get("license_date") or ""),
            "service_start_date": str(data.get("service_start_date") or ""),
            "media_class": media_class,
            "population_total": pop_total,
            "population_notes": pop_notes,
            "channels_count": len(smi_names),
            "rkn_url": rao_mod.build_rkn_url(lic_id),
            "org_name": data.get("org_name", ""),
            "smi_name": smi_name,
            "smi_names": smi_names,
            "smi_display": smi_display,
        })
    return items


def load_licenses_light(rkn_xlsx: Path, inn: str, include_radio: bool = False):
    return _licenses_light_cached(str(rkn_xlsx), rkn_xlsx.stat().st_mtime, inn, bool(include_radio))


_RKN_INDEX = None
_RKN_INDEX_MTIME = None
_RKN_INDEX_SRC = None
_RKN_INDEX_LOCK = threading.Lock()
_CALC_CONCURRENCY = max(1, int(os.getenv("RKN_CALC_CONCURRENCY", "1") or "1"))
_CALC_QUEUE_TIMEOUT_SECONDS = max(60, int(os.getenv("RKN_CALC_QUEUE_TIMEOUT_SECONDS", "600") or "600"))
_CALC_SEMAPHORE = threading.Semaphore(_CALC_CONCURRENCY)


def _build_rkn_index_from_sqlite(db_path: Path):
    idx = {}
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cols = {str(r[1]): int(r[0]) for r in cur.execute("PRAGMA table_info(rkn)")}
    has_activity = "licensed_activity" in cols
    has_license_date = "license_date" in cols
    has_service_start_date = "service_start_date" in cols
    has_freq = "freq" in cols
    has_brcst_descr = "brcst_descr" in cols
    sql = (
        "SELECT inn, org_name, license_num, "
        + ("license_date, " if has_license_date else "'' AS license_date, ")
        + ("service_start_date, " if has_service_start_date else "'' AS service_start_date, ")
        + ("licensed_activity, " if has_activity else "'' AS licensed_activity, ")
        + "sreda, population, status, smi_name14, smi_name, region_name_full, region_text, "
        + ("freq, " if has_freq else "'' AS freq, ")
        + ("brcst_descr " if has_brcst_descr else "'' AS brcst_descr ")
        + "FROM rkn"
    )
    for row in cur.execute(sql):
        inn, org_name, lic_id, license_date, service_start_date, licensed_activity, sreda, pop_raw, status, smi14, smi, region_name, region_text, freq, brcst_descr = row
        inn = (inn or "").strip()
        if not inn:
            continue
        if not _is_active_status(status):
            continue
        lic_id = (lic_id or "").strip()
        if not lic_id:
            continue
        by_lic = idx.setdefault(inn, {})
        lic = by_lic.setdefault(
            lic_id,
            {
                "org_name": (org_name or "").strip(),
                "media_raw": (sreda or "").strip(),
                "licensed_activity": (licensed_activity or "").strip(),
                "license_date": (license_date or "").strip(),
                "service_start_date": (service_start_date or "").strip(),
                "region_name_full": (region_name or "").strip(),
                "region_text": (region_text or "").strip(),
                "pop_values": set(),
                "pop_raw_values": set(),
                "pop_notes": [],
                "pop_rows": [],
                "smi_values": set(),
                "smi_display_values": set(),
                "smi_radio_flags": set(),
            },
        )
        if pop_raw is not None and str(pop_raw).strip() != "":
            lic["pop_raw_values"].add(str(pop_raw).strip())
            lic["pop_rows"].append({
                "population": pop_raw,
                "freq": freq,
                "brcst_descr": (brcst_descr or "").strip() or (sreda or "").strip(),
                "region_name_full": (region_name or "").strip(),
                "region_text": (region_text or "").strip(),
            })
        if not lic.get("licensed_activity") and licensed_activity:
            lic["licensed_activity"] = str(licensed_activity or "").strip()
        if not lic.get("license_date") and license_date:
            lic["license_date"] = str(license_date or "").strip()
        if not lic.get("service_start_date") and service_start_date:
            lic["service_start_date"] = str(service_start_date or "").strip()
        pop_int, pop_notes = rao_mod.parse_population(pop_raw)
        if pop_int is not None:
            lic["pop_values"].add(int(pop_int))
        if pop_notes:
            lic["pop_notes"].extend(pop_notes)
        _collect_smi_from_raw(lic, str(smi14 or "").strip(), str(smi or "").strip())
    conn.close()
    return idx


def _build_rkn_index_from_xlsx(rkn_xlsx: Path):
    header, it = _iter_rkn_rows_light(rkn_xlsx)
    idx = {h: i for i, h in enumerate(header)}

    def get(row, col):
        j = idx.get(col)
        if j is None:
            return None
        if j >= len(row):
            return None
        return row[j]

    out = {}
    for row in it:
        inn = str(get(row, "ns1:inn") or "").strip()
        if not inn:
            continue
        status = get(row, "ns1:status")
        if not _is_active_status(status):
            continue
        lic_id = str(get(row, "ns1:license_num") or "").strip()
        if not lic_id:
            continue
        org_name = str(get(row, "ns1:org_name") or "").strip()
        sreda = str(get(row, "ns1:sreda") or "").strip()
        pop_raw = get(row, "ns1:population")
        region_name = str(get(row, "ns1:region_name_full") or "").strip()
        region_text = str(get(row, "ns1:region_text") or "").strip()
        smi14 = str(get(row, "ns1:smi_name14") or "").strip()
        smi = str(get(row, "ns1:smi_name") or "").strip()
        licensed_activity = str(get(row, "ns1:licensed_activity") or "").strip()
        license_date = str(get(row, "ns1:license_date") or "").strip()
        service_start_date = str(get(row, "ns1:service_start_date") or "").strip()
        freq = get(row, "ns1:freq")
        brcst_descr = str(get(row, "ns1:brcst_descr") or "").strip()

        by_lic = out.setdefault(inn, {})
        lic = by_lic.setdefault(
            lic_id,
            {
                "org_name": org_name,
                "media_raw": sreda,
                "licensed_activity": licensed_activity,
                "license_date": license_date,
                "service_start_date": service_start_date,
                "region_name_full": region_name,
                "region_text": region_text,
                "pop_values": set(),
                "pop_raw_values": set(),
                "pop_notes": [],
                "pop_rows": [],
                "smi_values": set(),
                "smi_display_values": set(),
                "smi_radio_flags": set(),
            },
        )
        if pop_raw is not None and str(pop_raw).strip() != "":
            lic["pop_raw_values"].add(str(pop_raw).strip())
            lic["pop_rows"].append({
                "population": pop_raw,
                "freq": freq,
                "brcst_descr": brcst_descr or sreda,
                "region_name_full": region_name,
                "region_text": region_text,
            })
        if not lic.get("licensed_activity") and licensed_activity:
            lic["licensed_activity"] = licensed_activity
        if not lic.get("license_date") and license_date:
            lic["license_date"] = license_date
        if not lic.get("service_start_date") and service_start_date:
            lic["service_start_date"] = service_start_date
        pop_int, pop_notes = rao_mod.parse_population(pop_raw)
        if pop_int is not None:
            lic["pop_values"].add(int(pop_int))
        if pop_notes:
            lic["pop_notes"].extend(pop_notes)
        _collect_smi_from_raw(lic, smi14, smi)
    return out


def get_licenses_index():
    global _RKN_INDEX, _RKN_INDEX_MTIME, _RKN_INDEX_SRC
    with _RKN_INDEX_LOCK:
        db = find_rkn_db()
        if db:
            mtime = db.stat().st_mtime
            if (
                _RKN_INDEX is None
                or _RKN_INDEX_MTIME != mtime
                or _RKN_INDEX_SRC != str(db)
            ):
                _RKN_INDEX = _build_rkn_index_from_sqlite(db)
                _RKN_INDEX_MTIME = mtime
                _RKN_INDEX_SRC = str(db)
            return _RKN_INDEX

        rkn_xlsx = find_rkn_xlsx()
        mtime = rkn_xlsx.stat().st_mtime
        if (
            _RKN_INDEX is None
            or _RKN_INDEX_MTIME != mtime
            or _RKN_INDEX_SRC != str(rkn_xlsx)
        ):
            _RKN_INDEX = _build_rkn_index_from_xlsx(rkn_xlsx)
            _RKN_INDEX_MTIME = mtime
            _RKN_INDEX_SRC = str(rkn_xlsx)
        return _RKN_INDEX

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://ksenonya.github.io",
        "https://rnk-wof7.onrender.com",
        "http://localhost",
        "http://127.0.0.1:8000",
    ],
    allow_origin_regex=r"https://.*\\.github\\.io",
    allow_methods=["*"],
    allow_headers=["*"],
)

INN_MAP: dict[str, str] = {}
INN_MAP_MTIME: Optional[float] = None
INN_MAP_PATH: Optional[str] = None

def _resolve_inn_csv_path() -> Optional[Path]:
    env_path = os.getenv("RKN_INN_CSV_PATH", "").strip()
    if env_path:
        p = Path(env_path)
        if p.exists():
            return p
        # При явно заданном runtime-пути не переключаемся на bundled CSV:
        # иначе после web-upload можно получить новый XLSX/SQLite и старую карту ИНН.
        return None
    candidates = [
        TV_DATA_DIR / "inn_name.csv",
        BASE_DIR / "inn_name.csv",
        RKN_BOT_DIR / "inn_name.csv",
        Path.cwd() / "inn_name.csv",
    ]
    for p in candidates:
        if p.exists():
            return p
    return None


def _read_inn_csv(csv_path: Path) -> dict[str, str]:
    with csv_path.open("r", encoding="utf-8", newline="") as f:
        r = csv.DictReader(f)
        if not r.fieldnames:
            raise RuntimeError("inn_name.csv пустой или без заголовков")

        fields = set(r.fieldnames)
        inn_col = next((c for c in ("ns1:inn", "inn", "ИНН") if c in fields), None)
        name_col = next(
            (c for c in ("ns1:org_name_short", "ns1:org_name", "org_name_short", "org_name", "Наименование") if c in fields),
            None,
        )
        if not inn_col or not name_col:
            raise RuntimeError(
                f"inn_name.csv: не найдены колонки ИНН/наименования; есть {r.fieldnames}"
            )

        mp: dict[str, str] = {}
        for row in r:
            inn = (row.get(inn_col) or "").strip()
            name = (row.get(name_col) or "").strip()
            if not inn or not name:
                continue
            prev = mp.get(inn)
            if prev is None or (len(name), name) < (len(prev), prev):
                mp[inn] = name
        return mp


def refresh_inn_map_if_needed(force: bool = False) -> None:
    global INN_MAP, INN_MAP_MTIME, INN_MAP_PATH
    csv_path = _resolve_inn_csv_path()
    if not csv_path:
        if force:
            print("⚠️ inn_name.csv not found in calculators/rao_tv/data, backend dir or bot dir")
        return

    try:
        mtime = csv_path.stat().st_mtime
    except Exception:
        return

    if (
        not force
        and INN_MAP_PATH == str(csv_path)
        and INN_MAP_MTIME == mtime
        and INN_MAP
    ):
        return

    mp = _read_inn_csv(csv_path)
    INN_MAP = mp
    INN_MAP_MTIME = mtime
    INN_MAP_PATH = str(csv_path)
    print(f"✅ inn_name.csv loaded: {len(INN_MAP)} ({csv_path})")


def _org_name_from_rkn_index(inn: str) -> str:
    try:
        index = get_licenses_index()
        by_lic = index.get(inn, {})
        for data in by_lic.values():
            name = str(data.get("org_name") or "").strip()
            if name:
                return fix_mojibake(name)
    except Exception:
        return ""
    return ""


@app.on_event("startup")
def startup_warmup():
    try:
        _migrate_rkn_upload_storage_if_needed()
    except Exception:
        pass

    try:
        refresh_inn_map_if_needed(force=True)
    except Exception as e:
        print(f"⚠️ inn_name.csv load failed: {e}")

    try:
        idx = get_licenses_index()
        src = _RKN_INDEX_SRC or "unknown"
        size = len(idx or {})
        print(f"✅ RKN index warmed: {size} INN keys from {src}")
    except Exception as e:
        print(f"⚠️ RKN index warmup failed: {e}")

    # Прогрев статических таблиц ставок (снижает риск 504 на первом расчёте).
    try:
        vars_rao = find_vars_xlsx("rao_tv")
        vars_rao_mtime = float(vars_rao.stat().st_mtime)
        if hasattr(rao_mod, "_category_rate_map_cached"):
            rao_mod._category_rate_map_cached(str(vars_rao), vars_rao_mtime)
        if hasattr(rao_mod, "_topics_map_cached"):
            rao_mod._topics_map_cached(str(vars_rao), vars_rao_mtime)
        print(f"✅ rates cache warmed: {vars_rao}")
    except Exception as e:
        print(f"⚠️ rates cache warmup failed: {e}")

    if USAGE_AUTO_RECONCILE_ON_STARTUP:
        try:
            rec = _usage_reconcile_from_logs_monotonic(force=True)
            bf = _usage_backfill_auto_once()
            stats = _read_usage_stats()
            print(
                "✅ usage stats reconciled on startup: "
                f"total={int(stats.get('total') or 0)}, "
                f"items={len(stats.get('items') or [])}, "
                f"db={stats.get('db_path')}"
            )
            if rec:
                print(
                    "ℹ️ usage reconcile sources: "
                    f"event_files={len(rec.get('event_files') or [])}, "
                    f"generic_files={len(rec.get('generic_files') or [])}"
                )
            if bf:
                print(
                    "ℹ️ usage backfill scan: "
                    f"lines_scanned={int(bf.get('lines_scanned') or 0)}, "
                    f"matched={int(bf.get('lines_matched') or 0)}"
                )
        except Exception as e:
            print(f"⚠️ usage stats startup reconcile failed: {e}")



DASH_TOKENS = {"", "-", "—", "–", "нет"}


def _is_dash(v: Any) -> bool:
    if v is None:
        return True
    if isinstance(v, str):
        return v.strip().lower() in DASH_TOKENS
    return False


def _to_none_or_str(v: Any) -> Optional[str]:
    if _is_dash(v):
        return None
    return str(v).strip()


def _to_none_or_int(v: Any) -> Optional[int]:
    if _is_dash(v):
        return None
    if isinstance(v, int):
        return v
    if isinstance(v, float):
        return int(round(v))
    s = str(v).strip().replace("\u00a0", "").replace(" ", "")
    if not s:
        return None
    if s.isdigit() or (s.startswith("-") and s[1:].isdigit()):
        return int(s)
    if re.match(r"^-?\d+[.,]\d+$", s):
        return int(round(float(s.replace(",", "."))))
    compact = s.replace(",", "").replace(".", "")
    if compact.isdigit() or (compact.startswith("-") and compact[1:].isdigit()):
        return int(compact)
    raise ValueError(f"Неверное числовое значение: {v}")


def _clean_license_ids(v: Any) -> Optional[List[str]]:
    if v is None:
        return None
    if isinstance(v, str):
        parts = [p.strip() for p in v.split(",")]
        items = [p for p in parts if p]
    else:
        try:
            items = [str(x).strip() for x in list(v)]
        except Exception:
            items = []
    items = [x for x in items if x and not _is_dash(x)]
    return items or None


def _clean_population_by_license(v: Any) -> Optional[dict]:
    if v is None:
        return None
    if not isinstance(v, dict):
        return None
    out: dict[str, int] = {}
    for k, val in v.items():
        key = _to_none_or_str(k)
        if not key:
            continue
        n, _ = parse_population_override(val)
        if n is None:
            continue
        out[str(key)] = int(n)
    return out or None


def _clean_rate_by_license(v: Any) -> Optional[dict]:
    if v is None or not isinstance(v, dict):
        return None
    out: dict[str, float] = {}
    for k, val in v.items():
        key = _to_none_or_str(k)
        if not key:
            continue
        n = _to_none_or_float(val)
        if n is None or n < 0:
            continue
        out[str(key)] = float(n)
    return out or None


def _clean_int_by_license(v: Any) -> Optional[dict]:
    if v is None or not isinstance(v, dict):
        return None
    out: dict[str, int] = {}
    for k, val in v.items():
        key = _to_none_or_str(k)
        if not key:
            continue
        n = _to_none_or_int(val)
        if n is None or n < 0:
            continue
        out[str(key)] = int(n)
    return out or None


def _clean_text_by_license(v: Any) -> Optional[dict]:
    if v is None or not isinstance(v, dict):
        return None
    out: dict[str, str] = {}
    for k, val in v.items():
        key = _to_none_or_str(k)
        if not key:
            continue
        txt = _to_none_or_str(val)
        if not txt:
            continue
        out[str(key)] = str(txt).strip()
    return out or None


def _to_none_or_float(v: Any) -> Optional[float]:
    if _is_dash(v):
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip().replace(" ", "").replace(",", ".")
    if not s:
        return None
    return float(s)


class CalcRequest(BaseModel):
    inn: str = Field(..., description="ИНН 10 или 12 цифр")

    annual_revenue: Optional[float] = Field(None, ge=0)
    revenue_q: Optional[float] = Field(None, ge=0)
    expenses_q: Optional[float] = Field(None, ge=0)
    base_type: Optional[Literal["annual_revenue", "annual_expenses", "revenue_q", "expenses_q", "none"]] = None

    internet_resources: int = Field(0, ge=0, le=1000)
    contract_media: Literal["auto", "cable", "air", "both"] = "auto"

    new_user: bool = False
    assoc_member: bool = False
    fixed_fee_eligible: bool = False
    first_year_entity: bool = False
    manual_retransmission: bool = False

    only_license: Optional[str] = None
    license_ids: Optional[List[str]] = None
    new_license_ids: Optional[List[str]] = None
    population_override: Optional[int] = Field(None, ge=0, le=2_000_000_000)
    population_by_license: Optional[dict[str, int]] = None
    internet_resources_by_license: Optional[dict[str, int]] = None
    internet_sites_by_license: Optional[dict[str, str]] = None
    channel_founders: Optional[dict[str, str]] = None
    rate_by_license: Optional[dict[str, float]] = None
    actual_usage_share_by_channel: Optional[dict[str, float]] = None
    rate_mode: Literal["license", "actual_share"] = "license"
    actual_usage_share: Optional[float] = Field(None, ge=0, le=100)
    actual_usage_confirmed: bool = False
    actual_usage_change_ge20: bool = False
    actual_usage_periods: int = Field(0, ge=0, le=100)
    subscriber_total: Optional[int] = Field(None, ge=0, le=2_000_000_000)
    past_year_percent_paid: Optional[float] = Field(None, ge=0)
    calc_type: Literal["rao_tv", "vois_tv", "rao_radio", "vois_radio"] = "rao_tv"
    calc_mode: Literal["percent_minimum", "fixed"] = "percent_minimum"
    is_rar_member: bool = False
    is_state_budget_institution: bool = False
    is_internet_only: bool = False
    has_documented_income: bool = False
    is_100_state_capital: bool = False
    contract_period_number: int = Field(1, ge=1, le=200)
    special_min_period_number: int = Field(1, ge=1, le=200)
    is_new_user_and_other_use_contract: bool = False
    is_package_contract: bool = False
    has_legal_cases: bool = False
    signed_within_30_days: bool = True
    has_current_agreement_breach: bool = False
    previous_year_income_or_expense: Optional[float] = Field(None, ge=0)
    fixed_fee_basis_confirmed: bool = False
    market_index: float = Field(1.0, ge=0)
    inflation_index: float = Field(1.0, ge=0)
    simultaneous_internet_broadcast: bool = False
    internet_admin_by_user: bool = False
    site_app_count: int = Field(0, ge=0, le=1000)
    has_third_party_channels: bool = False

    if _V2:
        @field_validator("inn", mode="before")
        @classmethod
        def _v_inn(cls, v: Any) -> str:
            s = _to_none_or_str(v)
            if not s:
                raise ValueError("ИНН обязателен")
            s = s.replace(" ", "")
            if not s.isdigit() or len(s) not in (10, 12):
                raise ValueError("ИНН должен состоять из 10 или 12 цифр")
            return s

        @field_validator(
            "annual_revenue",
            "revenue_q",
            "expenses_q",
            "past_year_percent_paid",
            "actual_usage_share",
            "previous_year_income_or_expense",
            "market_index",
            "inflation_index",
            mode="before"
        )
        @classmethod
        def _v_floats(cls, v: Any) -> Optional[float]:
            return _to_none_or_float(v)

        @field_validator("only_license", mode="before")
        @classmethod
        def _v_only_license(cls, v: Any) -> Optional[str]:
            return _to_none_or_str(v)

        @field_validator("license_ids", mode="before")
        @classmethod
        def _v_license_ids(cls, v: Any) -> Optional[List[str]]:
            return _clean_license_ids(v)

        @field_validator("new_license_ids", mode="before")
        @classmethod
        def _v_new_license_ids(cls, v: Any) -> Optional[List[str]]:
            return _clean_license_ids(v)

        @field_validator("population_override", mode="before")
        @classmethod
        def _v_pop(cls, v: Any) -> Optional[int]:
            return _to_none_or_int(v)

        @field_validator("population_by_license", mode="before")
        @classmethod
        def _v_pop_by_license(cls, v: Any) -> Optional[dict]:
            return _clean_population_by_license(v)

        @field_validator("internet_resources_by_license", mode="before")
        @classmethod
        def _v_internet_by_license(cls, v: Any) -> Optional[dict]:
            return _clean_int_by_license(v)

        @field_validator("internet_sites_by_license", mode="before")
        @classmethod
        def _v_internet_sites_by_license(cls, v: Any) -> Optional[dict]:
            return _clean_text_by_license(v)

        @field_validator("rate_by_license", mode="before")
        @classmethod
        def _v_rate_by_license(cls, v: Any) -> Optional[dict]:
            return _clean_rate_by_license(v)

        @field_validator("actual_usage_share_by_channel", mode="before")
        @classmethod
        def _v_actual_share_by_channel(cls, v: Any) -> Optional[dict]:
            return _clean_rate_by_license(v)

    else:
        @field_validator("inn", pre=True)
        def _v1_inn(cls, v: Any) -> str:
            s = _to_none_or_str(v)
            if not s:
                raise ValueError("ИНН обязателен")
            s = s.replace(" ", "")
            if not s.isdigit() or len(s) not in (10, 12):
                raise ValueError("ИНН должен состоять из 10 или 12 цифр")
            return s

        @field_validator(
            "annual_revenue",
            "revenue_q",
            "expenses_q",
            "past_year_percent_paid",
            "actual_usage_share",
            "previous_year_income_or_expense",
            "market_index",
            "inflation_index",
            pre=True
        )
        def _v1_floats(cls, v: Any) -> Optional[float]:
            return _to_none_or_float(v)

        @field_validator("only_license", pre=True)
        def _v1_only_license(cls, v: Any) -> Optional[str]:
            return _to_none_or_str(v)

        @field_validator("license_ids", pre=True)
        def _v1_license_ids(cls, v: Any) -> Optional[List[str]]:
            return _clean_license_ids(v)

        @field_validator("new_license_ids", pre=True)
        def _v1_new_license_ids(cls, v: Any) -> Optional[List[str]]:
            return _clean_license_ids(v)

        @field_validator("population_override", pre=True)
        def _v1_pop(cls, v: Any) -> Optional[int]:
            return _to_none_or_int(v)

        @field_validator("population_by_license", pre=True)
        def _v1_pop_by_license(cls, v: Any) -> Optional[dict]:
            return _clean_population_by_license(v)

        @field_validator("internet_resources_by_license", pre=True)
        def _v1_internet_by_license(cls, v: Any) -> Optional[dict]:
            return _clean_int_by_license(v)

        @field_validator("internet_sites_by_license", pre=True)
        def _v1_internet_sites_by_license(cls, v: Any) -> Optional[dict]:
            return _clean_text_by_license(v)

        @field_validator("rate_by_license", pre=True)
        def _v1_rate_by_license(cls, v: Any) -> Optional[dict]:
            return _clean_rate_by_license(v)

        @field_validator("actual_usage_share_by_channel", pre=True)
        def _v1_actual_share_by_channel(cls, v: Any) -> Optional[dict]:
            return _clean_rate_by_license(v)


class ExportDocxRequest(BaseModel):
    title: str = Field(..., min_length=1, max_length=300)
    text: str = Field(..., min_length=1)
    html: Optional[str] = None
    summary: Optional[List[dict[str, str]]] = None


class UsageStatsBackfillRequest(BaseModel):
    login: str
    password: str
    log_globs: Optional[list[str]] = None
    dry_run: bool = False


class EventCalcRequest(BaseModel):
    calc_type: Literal["rao_events", "vois_events"] = "rao_events"
    user_type: Literal["standard", "special"]
    special_calculation_type: Optional[Literal["quarterly", "hourly"]] = None
    has_reliable_data: bool = True
    visitors: Optional[int] = Field(None, ge=0)
    duration_hours: Optional[float] = Field(None, gt=0)
    duration_hours_special: Optional[float] = Field(None, gt=0)
    region: Optional[str] = None
    is_free_access: bool = False
    music_share: Literal["19", "39", "59", "79", "100", "unknown"] = "100"
    admin_platforms_count: Optional[int] = Field(None, ge=0)
    external_platforms_count: Optional[int] = Field(None, ge=0)
    broadcasts_count: Optional[int] = Field(None, ge=1)
    quarterly_contract_year_confirmed: bool = True
    quarterly_prepay_4q_confirmed: bool = True

    # Обратная совместимость со старым фронтендом:
    has_admin_platforms: Optional[bool] = None
    external_platforms: Optional[int] = Field(None, ge=0)

@app.get("/", response_class=HTMLResponse)
def home():
    headers = {
        "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
        "Pragma": "no-cache",
        "Expires": "0",
    }
    if not INDEX_HTML.exists():
        return HTMLResponse(
            "<h1>index.html не найден</h1><p>Проверь наличие /interface/index.html</p>",
            status_code=500,
            headers=headers,
        )
    return HTMLResponse(INDEX_HTML.read_text(encoding="utf-8"), headers=headers)


@app.get("/favicon.svg")
def favicon_svg():
    headers = {
        "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
        "Pragma": "no-cache",
        "Expires": "0",
    }
    if not FAVICON_SVG.exists():
        return Response(status_code=404, headers=headers)
    return FileResponse(FAVICON_SVG, media_type="image/svg+xml", headers=headers)


@app.get("/favicon.ico")
def favicon_ico():
    headers = {
        "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
        "Pragma": "no-cache",
        "Expires": "0",
    }
    if not FAVICON_ICO.exists():
        return Response(status_code=404, headers=headers)
    return FileResponse(FAVICON_ICO, media_type="image/x-icon", headers=headers)


@app.get("/favicon-32x32.png")
def favicon_png():
    headers = {
        "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
        "Pragma": "no-cache",
        "Expires": "0",
    }
    if not FAVICON_PNG.exists():
        return Response(status_code=404, headers=headers)
    return FileResponse(FAVICON_PNG, media_type="image/png", headers=headers)


@app.get("/apple-touch-icon.png")
def apple_touch_icon():
    headers = {
        "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
        "Pragma": "no-cache",
        "Expires": "0",
    }
    if not APPLE_TOUCH_ICON.exists():
        return Response(status_code=404, headers=headers)
    return FileResponse(APPLE_TOUCH_ICON, media_type="image/png", headers=headers)




@app.get("/api/inninfo")
def api_inninfo(inn: str):
    try:
        inn_clean = parse_inn(inn)
    except Exception as e:
        return JSONResponse(status_code=400, content={"ok": False, "error": str(e)})

    try:
        refresh_inn_map_if_needed()
    except Exception:
        pass

    org_name = INN_MAP.get(inn_clean)
    if org_name:
        org_name = fix_mojibake(org_name)
        return {"ok": True, "inn": inn_clean, "org_name": org_name}
    return JSONResponse(status_code=404, content={"ok": False, "org_name": ""})


@app.get("/api/licenses")
def api_licenses(inn: str, include_radio: bool = False, calc_type: str = "rao_tv"):
    try:
        inn_clean = parse_inn(inn)
    except Exception as e:
        return JSONResponse(status_code=400, content={"ok": False, "error": str(e)})

    mode = str(calc_type or "").strip().lower()
    radio_only = mode in {"rao_radio", "vois_radio", "vois_rv"}

    def _collect_items(by_lic: dict[str, dict]) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        for lic_id, data in by_lic.items():
            flags = data.get("smi_radio_flags") or set()
            all_radio = bool(flags) and all(bool(x) for x in flags)
            media_raw = str(data.get("media_raw") or "")
            licensed_activity = str(data.get("licensed_activity") or "")
            activity_is_radio = _activity_is_radio(licensed_activity)
            activity_is_tv = _activity_is_tv(licensed_activity)
            media_is_radio = ("радио" in media_raw.lower()) or ("радиоканал" in media_raw.lower()) or ("радиовещ" in media_raw.lower())
            if (not radio_only) and (not include_radio) and (all_radio or activity_is_radio):
                continue
            media_l = media_raw.lower()
            media_is_tv = ("теле" in media_l) or ("телеканал" in media_l) or ("телевещ" in media_l)
            has_radio_flag = any(bool(x) for x in flags)
            if radio_only:
                # licensed_activity — приоритетный сигнал.
                # Далее используем media/smi как fallback для старых выгрузок.
                if activity_is_tv and not activity_is_radio:
                    continue
                if (
                    not activity_is_radio
                    and (media_is_tv and not media_is_radio and not has_radio_flag)
                ):
                    continue
            pop_total, pop_notes = _normalize_license_population_fast(data)
            if radio_only:
                pop_total, pop_notes = _cap_radio_population(pop_total, pop_notes)
            media_raw = data.get("media_raw") or ""
            media_class = rao_mod.normalize_media(media_raw)
            smi_names = _sorted_unique(data.get("smi_values"))
            smi_display_list = _sorted_unique(data.get("smi_display_values"))
            smi_name = smi_names[0] if smi_names else ""
            smi_display = _format_smi_display(smi_display_list) if smi_display_list else smi_name
            items.append({
                "license_id": lic_id,
                "media_raw": media_raw,
                "licensed_activity": licensed_activity,
                "license_date": str(data.get("license_date") or ""),
                "service_start_date": str(data.get("service_start_date") or ""),
                "media_class": media_class,
                "population_total": pop_total,
                "population_notes": pop_notes,
                "channels_count": len(smi_names),
                "rkn_url": rao_mod.build_rkn_url(lic_id),
                "org_name": data.get("org_name", ""),
                "smi_name": smi_name,
                "smi_names": smi_names,
                "smi_display": smi_display,
            })
        return items

    try:
        index = get_licenses_index()
        by_lic = index.get(inn_clean, {})
        items = _collect_items(by_lic)
        return {"ok": True, "inn": inn_clean, "licenses": items, "notes": []}
    except Exception as e:
        try:
            # Если индекс уже был собран ранее — используем его как "последнюю стабильную версию"
            # и не уходим в тяжёлый XLSX-fallback, который может приводить к 504.
            stale = _RKN_INDEX if isinstance(_RKN_INDEX, dict) else {}
            if stale:
                by_lic = stale.get(inn_clean, {})
                items = _collect_items(by_lic)
                return {"ok": True, "inn": inn_clean, "licenses": items, "notes": ["Использована последняя стабильная версия индекса лицензий."]}
        except Exception:
            pass
        return JSONResponse(
            status_code=503,
            content={"ok": False, "error": "Сервис лицензий временно занят обновлением данных. Повторите попытку через 10–20 секунд."},
        )


@app.post("/api/calc")
@app.post("/api/calc/")
def api_calc(req: CalcRequest):
    argv: List[str] = ["--inn", req.inn.strip(), "--non_interactive"]
    calc_type = (req.calc_type or "rao_tv").strip().lower()

    if calc_type in {"rao_radio", "vois_radio", "vois_rv"}:
        radio_mode = "vois_radio" if calc_type in {"vois_radio", "vois_rv"} else "rao_radio"
        rargv: List[str] = ["--inn", req.inn.strip(), "--non_interactive"]
        try:
            rargv += ["--rkn_xlsx", str(find_rkn_xlsx())]
        except Exception as e:
            return JSONResponse(status_code=500, content={"ok": False, "error": f"Ошибка загрузки таблицы РКН: {e}"})
        try:
            radio_vars_type = "vois_tv" if radio_mode == "vois_radio" else "rao_tv"
            rargv += ["--vars_xlsx", str(find_vars_xlsx(radio_vars_type))]
        except Exception as e:
            return JSONResponse(status_code=500, content={"ok": False, "error": f"Ошибка загрузки таблицы ставок ({radio_mode}): {e}"})

        if req.revenue_q is not None:
            rargv += ["--revenue_q", str(req.revenue_q)]
        if req.annual_revenue is not None:
            rargv += ["--annual_revenue", str(req.annual_revenue)]
        if req.expenses_q is not None:
            rargv += ["--expenses_q", str(req.expenses_q)]
        if req.base_type:
            rargv += ["--base_type", str(req.base_type)]
        if req.license_ids:
            for lic in req.license_ids:
                if lic:
                    rargv += ["--licenses", str(lic).strip()]
        if req.new_license_ids:
            for lic in req.new_license_ids:
                if lic:
                    rargv += ["--new_license_ids", str(lic).strip()]
        if req.calc_mode:
            rargv += ["--calc_mode", str(req.calc_mode)]
        if req.is_rar_member:
            rargv.append("--is_rar_member")
        if req.is_state_budget_institution:
            rargv.append("--is_state_budget_institution")
        if req.is_internet_only:
            rargv.append("--is_internet_only")
        if req.has_documented_income:
            rargv.append("--has_documented_income")
        if req.is_100_state_capital:
            rargv.append("--is_100_state_capital")
        if req.new_user:
            rargv.append("--new_user")
        if req.contract_period_number:
            rargv += ["--contract_period_number", str(int(req.contract_period_number))]
        if req.special_min_period_number:
            rargv += ["--special_min_period_number", str(int(req.special_min_period_number))]
        if req.is_new_user_and_other_use_contract:
            rargv.append("--is_new_user_and_other_use_contract")
        if req.is_package_contract:
            rargv.append("--is_package_contract")
        if req.assoc_member:
            rargv.append("--assoc_member")
        if req.has_legal_cases:
            rargv.append("--has_legal_cases")
        if req.signed_within_30_days:
            rargv.append("--signed_within_30_days")
        if req.has_current_agreement_breach:
            rargv.append("--has_current_agreement_breach")
        if req.previous_year_income_or_expense is not None:
            rargv += ["--previous_year_income_or_expense", str(float(req.previous_year_income_or_expense))]
        if req.fixed_fee_eligible:
            rargv.append("--fixed_fee_eligible")
        if req.fixed_fee_basis_confirmed:
            rargv.append("--fixed_fee_basis_confirmed")
        if req.market_index is not None:
            rargv += ["--market_index", str(float(req.market_index))]
        if req.inflation_index is not None:
            rargv += ["--inflation_index", str(float(req.inflation_index))]
        if req.simultaneous_internet_broadcast:
            rargv.append("--simultaneous_internet_broadcast")
        if req.internet_admin_by_user:
            rargv.append("--internet_admin_by_user")
        if req.site_app_count:
            rargv += ["--site_app_count", str(int(req.site_app_count))]
        if req.internet_resources_by_license:
            for lic, cnt in req.internet_resources_by_license.items():
                rargv += ["--internet_resources_by_license", f"{lic}={int(cnt)}"]
        if req.internet_sites_by_license:
            for lic, raw_sites in req.internet_sites_by_license.items():
                sites = [x.strip() for x in re.split(r"[,\n;|]+", str(raw_sites or "")) if x.strip()]
                if sites:
                    rargv += ["--internet_sites_by_license", f"{lic}={'|'.join(sites)}"]
        if req.channel_founders:
            for key, founder in req.channel_founders.items():
                value = str(founder or "").strip().lower()
                if value in {"user", "third_party"}:
                    rargv += ["--channel_founder", f"{key}={value}"]
        if req.rate_mode == "actual_share" and req.actual_usage_share_by_channel:
            for key, share in req.actual_usage_share_by_channel.items():
                rargv += ["--actual_share_by_channel", f"{key}={float(share)}"]
        if req.has_third_party_channels:
            rargv.append("--has_third_party_channels")
        rargv += ["--society", "ВОИС" if radio_mode == "vois_radio" else "РАО"]
        rargv.append("--skip_runtime_population_normalization")

        acquired = _CALC_SEMAPHORE.acquire(timeout=_CALC_QUEUE_TIMEOUT_SECONDS)
        if not acquired:
            return JSONResponse(
                status_code=503,
                content={"ok": False, "error": "Сервис расчёта перегружен. Повторите попытку через 10-20 секунд."},
            )
        try:
            code, out = run_radio_calc_capture(rargv)
            out = (out or "").strip()
        finally:
            _CALC_SEMAPHORE.release()

        if code == 0:
            _increment_usage(radio_mode)
            if out.lstrip().startswith("<article"):
                radio_model = extract_radio_report_model_from_html(out)
                plain_text = render_radio_report_text(radio_model) if radio_model is not None else _report_html_to_text(out)
                return JSONResponse(status_code=200, content={
                    "ok": True,
                    "text": plain_text,
                    "html": out,
                })
            return JSONResponse(status_code=200, content={"ok": True, "text": out})
        return JSONResponse(status_code=400, content={"ok": False, "error": out or "Ошибка расчёта"})

    try:
        argv += ["--rkn_xlsx", str(find_rkn_xlsx())]
    except Exception as e:
        return JSONResponse(status_code=500, content={"ok": False, "error": f"Ошибка загрузки таблицы РКН: {e}"})
    try:
        argv += ["--vars_xlsx", str(find_vars_xlsx(calc_type))]
    except Exception as e:
        return JSONResponse(status_code=500, content={"ok": False, "error": f"Ошибка загрузки таблицы ставок ({calc_type}): {e}"})
    argv += ["--society", "ВОИС" if calc_type == "vois_tv" else "РАО"]

    if req.revenue_q is not None:
        argv += ["--revenue_q", str(req.revenue_q)]
    if req.annual_revenue is not None:
        argv += ["--annual_revenue", str(req.annual_revenue)]
    if req.expenses_q is not None:
        argv += ["--expenses_q", str(req.expenses_q)]

    argv += ["--internet_resources", str(req.internet_resources)]
    argv += ["--contract_quarter", str(int(req.contract_period_number or 1))]
    argv += ["--contract_media", str(req.contract_media)]

    if req.new_user:
        argv.append("--new_user")
    if req.assoc_member:
        argv.append("--assoc_member")
    if req.is_package_contract:
        argv.append("--is_package_contract")
    if req.is_new_user_and_other_use_contract:
        argv.append("--is_new_user_and_other_use_contract")
    if req.fixed_fee_eligible:
        argv.append("--fixed_fee_eligible")
    if req.first_year_entity:
        argv.append("--first_year_entity")
    if req.manual_retransmission:
        argv.append("--manual_retransmission")

    if req.only_license:
        argv += ["--only_license", req.only_license.strip()]
    if req.license_ids:
        for lic in req.license_ids:
            if lic:
                argv += ["--licenses", str(lic).strip()]
    if req.population_override is not None:
        argv += ["--population_override", str(int(req.population_override))]
    if req.population_by_license:
        for lic, pop in req.population_by_license.items():
            argv += ["--population_by_license", f"{lic}={int(pop)}"]
    if req.internet_resources_by_license:
        for lic, cnt in req.internet_resources_by_license.items():
            argv += ["--internet_resources_by_license", f"{lic}={int(cnt)}"]
    if req.internet_sites_by_license:
        for lic, raw_sites in req.internet_sites_by_license.items():
            sites = [x.strip() for x in re.split(r"[,\n;|]+", str(raw_sites or "")) if x.strip()]
            if sites:
                argv += ["--internet_sites_by_license", f"{lic}={'|'.join(sites)}"]
    if req.rate_by_license:
        for lic, rate in req.rate_by_license.items():
            argv += ["--rate_by_license", f"{lic}={float(rate)}"]
    argv += ["--rate_mode", str(req.rate_mode or "license")]
    if req.rate_mode == "actual_share" and req.actual_usage_share_by_channel:
        for key, share in req.actual_usage_share_by_channel.items():
            argv += ["--actual_share_by_channel", f"{key}={float(share)}"]
    if req.actual_usage_share is not None:
        argv += ["--actual_usage_share", str(float(req.actual_usage_share))]
    if req.actual_usage_confirmed:
        argv.append("--actual_usage_confirmed")
    if req.actual_usage_change_ge20:
        argv.append("--actual_usage_change_ge20")
    if req.actual_usage_periods:
        argv += ["--actual_usage_periods", str(int(req.actual_usage_periods))]
    if req.subscriber_total is not None:
        argv += ["--subscriber_total", str(int(req.subscriber_total))]
    if req.past_year_percent_paid is not None:
        argv += ["--past_year_percent_paid", str(float(req.past_year_percent_paid))]

    # Расчёт всегда выполняем без runtime-сетевой нормализации населения.
    # Нормализация должна происходить на шаге выбора лицензий.
    argv.append("--skip_runtime_population_normalization")

    acquired = _CALC_SEMAPHORE.acquire(timeout=_CALC_QUEUE_TIMEOUT_SECONDS)
    if not acquired:
        return JSONResponse(
            status_code=503,
            content={"ok": False, "error": "Сервис расчёта перегружен. Повторите попытку через 10-20 секунд."},
        )
    try:
        code, out = run_calc_capture(argv)
        out = (out or "").strip()
    finally:
        _CALC_SEMAPHORE.release()

    if code == 0:
        _increment_usage(calc_type)
        title = (
            "Калькулятор ставок ВОИС для вещателей телеканалов"
            if calc_type == "vois_tv"
            else "Калькулятор ставок РАО для вещателей телеканалов"
        )
        return JSONResponse(status_code=200, content={
            "ok": True,
            "text": out,
            "html": _plain_result_to_report_html(out, title),
        })
    return JSONResponse(status_code=400, content={"ok": False, "error": out or "Ошибка расчёта"})


@app.post("/api/events/calc")
@app.post("/api/events/calc/")
def api_events_calc(req: EventCalcRequest):
    try:
        admin_count = req.admin_platforms_count
        if admin_count is None:
            admin_count = 1 if bool(req.has_admin_platforms) else 0

        external_count = req.external_platforms_count
        if external_count is None:
            external_count = int(req.external_platforms or 0)

        payload = {
            "user_type": req.user_type,
            "special_calculation_type": req.special_calculation_type,
            "has_reliable_data": req.has_reliable_data,
            "visitors": req.visitors,
            "duration_hours": req.duration_hours,
            "duration_hours_special": req.duration_hours_special,
            "region": req.region,
            "is_free_access": req.is_free_access,
            "music_share": req.music_share,
            "admin_platforms_count": max(0, int(admin_count)),
            "external_platforms_count": max(0, int(external_count)),
            "broadcasts_count": max(1, int(req.broadcasts_count or 1)),
            "quarterly_contract_year_confirmed": bool(req.quarterly_contract_year_confirmed),
            "quarterly_prepay_4q_confirmed": bool(req.quarterly_prepay_4q_confirmed),
        }

        if req.calc_type == "vois_events":
            result = calculate_vois_event_fee(payload)
            _increment_usage("vois_events")
        else:
            result = calculate_rao_event_fee(payload)
            _increment_usage("rao_events")
        text = str(result.get("text") or "").strip()
        provider = "ВОИС" if req.calc_type == "vois_events" else "РАО"
        title = (
            "Калькулятор ставок ВОИС для трансляторов мероприятий"
            if req.calc_type == "vois_events"
            else "Калькулятор ставок РАО для трансляторов мероприятий"
        )
        if text:
            result["html"] = _event_result_to_report_html(
                result,
                payload,
                provider=provider,
                title=title,
            )
        return JSONResponse(status_code=200, content=result)
    except ValueError as e:
        return JSONResponse(status_code=400, content={"ok": False, "error": str(e)})
    except Exception as e:
        return JSONResponse(status_code=500, content={"ok": False, "error": f"Ошибка расчёта онлайн-трансляции: {e}"})


@app.post("/api/export-docx")
@app.post("/api/export-docx/")
def api_export_docx(req: ExportDocxRequest):
    try:
        from docx import Document  # type: ignore
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore
        from docx.shared import Inches, Pt, RGBColor  # type: ignore
    except Exception:
        return JSONResponse(status_code=500, content={"ok": False, "error": "На сервере не установлена библиотека python-docx"})

    title = (req.title or "").strip()
    text = (req.text or "").strip()
    if not title or not text:
        return JSONResponse(status_code=400, content={"ok": False, "error": "Для экспорта нужен заголовок и текст отчёта"})

    report_html = str(req.html or "").strip()
    radio_model = extract_radio_report_model_from_html(report_html)
    if radio_model is not None:
        marker = "Результат расчёта по настоящему калькулятору"
        disclaimer_pos = text.find(marker)
        disclaimer = text[disclaimer_pos:].strip() if disclaimer_pos >= 0 else ""
        content = build_radio_report_docx(radio_model, disclaimer=disclaimer)
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="rnk-calculation.docx"'},
        )

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.bold = False
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Heading 1", 18, RGBColor(20, 33, 52)),
        ("Heading 2", 13, RGBColor(32, 76, 114)),
        ("Heading 3", 11, RGBColor(32, 76, 114)),
    ]:
        try:
            style = styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = color
            style.paragraph_format.space_before = Pt(8)
            style.paragraph_format.space_after = Pt(5)
        except Exception:
            pass

    def _set_cell_shading(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def _set_cell_width(cell, width: int) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(width))
        tc_w.set(qn("w:type"), "dxa")

    def _set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
            node = tc_mar.find(qn(f"w:{name}"))
            if node is None:
                node = OxmlElement(f"w:{name}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def _set_table_borders(table, color: str = "D7E2EF") -> None:
        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge = borders.find(qn(f"w:{name}"))
            if edge is None:
                edge = OxmlElement(f"w:{name}")
                borders.append(edge)
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), "6")
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), color)

    def _format_cell(cell, *, shaded: bool = False, fill: str = "F7FAFC", bold: bool = False) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if shaded:
            _set_cell_shading(cell, fill)
        _set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.bold = bold
                run.font.color.rgb = RGBColor(20, 33, 52)

    def _add_heading(value: str, level: int = 2):
        p = doc.add_heading(str(value or "").strip(), level=level)
        return p

    def _add_text(
        value: str,
        *,
        bold: bool = False,
        italic: bool = False,
        size: int = 10,
        color: RGBColor | None = None,
        highlight: bool = False,
    ):
        p = doc.add_paragraph()
        run = p.add_run(str(value or "").strip())
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color or RGBColor(20, 33, 52)
        if highlight:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return p

    def _add_hyperlink_text(
        value: str,
        href: str,
        *,
        bold: bool = False,
        size: int = 10,
        color: RGBColor | None = None,
    ):
        p = doc.add_paragraph()
        text_value = str(value or "").strip()
        link_value = str(href or "").strip()
        if not link_value:
            run = p.add_run(text_value)
            run.font.name = "Arial"
            run.font.size = Pt(size)
            run.bold = bold
            run.font.color.rgb = color or RGBColor(20, 33, 52)
            return p
        try:
            r_id = doc.part.relate_to(link_value, RT.HYPERLINK, is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)
            new_run = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            if bold:
                r_pr.append(OxmlElement("w:b"))
            color_node = OxmlElement("w:color")
            color_node.set(qn("w:val"), "204C72")
            r_pr.append(color_node)
            underline = OxmlElement("w:u")
            underline.set(qn("w:val"), "single")
            r_pr.append(underline)
            new_run.append(r_pr)
            text_node = OxmlElement("w:t")
            text_node.text = text_value
            new_run.append(text_node)
            hyperlink.append(new_run)
            p._p.append(hyperlink)
        except Exception:
            run = p.add_run(text_value)
            run.font.name = "Arial"
            run.font.size = Pt(size)
            run.bold = bold
            run.font.color.rgb = color or RGBColor(32, 76, 114)
        return p

    def _strip_html_fragment(value: Any) -> str:
        raw = str(value or "")
        raw = re.sub(r"(?is)<br\s*/?>", "\n", raw)
        raw = re.sub(
            r'(?is)<a\b[^>]*href=["\']([^"\']+)["\'][^>]*>(.*?)</a>',
            lambda m: _strip_html_fragment(m.group(2)),
            raw,
        )
        raw = re.sub(r"(?is)<[^>]+>", "", raw)
        raw = html.unescape(raw)
        raw = re.sub(r"[ \t]+", " ", raw)
        raw = re.sub(r"\n\s+", "\n", raw)
        return raw.strip()

    def _extract_first_href(value: Any) -> str:
        match = re.search(r'(?is)<a\b[^>]*href=["\']([^"\']+)["\']', str(value or ""))
        return html.unescape(match.group(1)).strip() if match else ""

    def _report_nodes_from_html(raw_html: str) -> list[dict[str, str]]:
        nodes: list[dict[str, str]] = []
        raw_html = re.sub(
            r'(?is)<details\b[^>]*class=["\'][^"\']*eventReportDetails[^"\']*["\'][^>]*>.*?</details>',
            "",
            str(raw_html or ""),
        )

        def _append_node(kind: str, content: str, value: str = "") -> None:
            href = _extract_first_href(content) or _extract_first_href(value)
            content = _strip_html_fragment(content)
            value = _strip_html_fragment(value)
            if kind == "row":
                content = re.sub(r":\s*$", "", content)
            if content or value:
                nodes.append({"kind": kind, "text": content, "value": value, "href": href})

        for token in re.finditer(
            r'(?is)<h1[^>]*>(.*?)</h1>|'
            r'<h2[^>]*>(.*?)</h2>|'
            r'<h3[^>]*>(.*?)</h3>|'
            r'<h4[^>]*>(.*?)</h4>|'
            r'<(?:p|div) class="([^"]*radioReportFact[^"]*)">\s*<strong>(.*?):?</strong>\s*(.*?)</(?:p|div)>|'
            r'<div class="([^"]*radioReportRow[^"]*)">\s*<div class="radioReportKey">(.*?)</div>\s*<div class="radioReportValue">(.*?)</div>\s*</div>|'
            r'<div class="[^"]*eventReportTotal[^"]*">(.*?)</div>|'
            r'<div class="[^"]*eventReportLine[^"]*">\s*<strong>(.*?):?</strong>\s*<span>(.*?)</span>\s*</div>|'
            r'<div class="[^"]*radioReportLicenseTitle[^"]*">(.*?)</div>|'
            r'<div class="[^"]*radioReportLicenseAddon[^"]*">(.*?)</div>|'
            r'<div class="radioReportChannelTitle">(.*?)</div>|'
            r'<li[^>]*>(.*?)</li>|'
            r'<p[^>]*>(.*?)</p>',
            raw_html,
        ):
            (
                h1,
                h2,
                h3,
                h4,
                fact_class,
                fact_key,
                fact_value,
                row_class,
                key,
                value,
                event_total,
                event_key,
                event_value,
                lic,
                addon,
                channel,
                li,
                p,
            ) = token.groups()
            if h1 is not None:
                _append_node("h1", h1)
            elif h2 is not None:
                _append_node("h2", h2)
            elif h3 is not None:
                _append_node("h3", h3)
            elif h4 is not None:
                _append_node("h4", h4)
            elif fact_key is not None:
                kind = "result_row" if "radioReportResultFact" in (fact_class or "") else "row"
                _append_node(kind, fact_key, fact_value or "")
            elif key is not None:
                kind = "result_row" if "radioReportResultRow" in (row_class or "") else "row"
                _append_node(kind, key, value or "")
            elif event_total is not None:
                total_text = _strip_html_fragment(event_total)
                if ":" in total_text:
                    total_key, total_value = total_text.split(":", 1)
                    _append_node("result_row", total_key, total_value)
                else:
                    _append_node("result_text", event_total)
            elif event_key is not None:
                _append_node("row", event_key, event_value or "")
            elif lic is not None:
                _append_node("license", lic)
            elif addon is not None:
                _append_node("addon", addon)
            elif channel is not None:
                _append_node("channel", channel)
            elif li is not None:
                _append_node("li", li)
            elif p is not None:
                _append_node("p", p)
        return nodes

    def _add_rows_table(rows: list[tuple[str, str, bool]]) -> None:
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        table.autofit = False
        _set_table_borders(table)
        try:
            table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:
            pass
        for i, (key, value, is_result) in enumerate(rows):
            left = table.cell(i, 0)
            right = table.cell(i, 1)
            left.text = key
            right.text = value
            _set_cell_width(left, 3000)
            _set_cell_width(right, 6000)
            if is_result:
                _format_cell(left, shaded=True, fill="EEF6F2", bold=True)
                _format_cell(right, shaded=True, fill="EEF6F2", bold=True)
            else:
                _format_cell(left, shaded=True, fill="F7FAFC", bold=True)
                _format_cell(right)

    def _add_report_nodes(nodes: list[dict[str, str]]) -> None:
        pending_rows: list[tuple[str, str, bool]] = []

        def flush_rows() -> None:
            nonlocal pending_rows
            _add_rows_table(pending_rows)
            pending_rows = []

        def is_disclaimer(value: str) -> bool:
            return str(value or "").strip().startswith("Результат расчёта по настоящему калькулятору")

        for node in nodes:
            kind = node.get("kind")
            value = node.get("text", "")
            node_value = node.get("value", "")
            href = node.get("href", "")
            if is_disclaimer(value) or is_disclaimer(node_value):
                continue
            if kind in {"row", "result_row"}:
                pending_rows.append((value, node_value, kind == "result_row"))
                continue
            flush_rows()
            if kind == "h1":
                if value.strip().lower() == "результат расчёта":
                    continue
                _add_heading(value, level=2)
            elif kind == "h2":
                if value.strip().lower() == "результат расчёта":
                    continue
                _add_heading(value, level=2)
            elif kind == "h3":
                _add_heading(value, level=3)
            elif kind == "h4":
                _add_text(value, bold=True, size=10, color=RGBColor(32, 76, 114))
            elif kind == "license":
                _add_hyperlink_text(value, href, bold=True, size=10, color=RGBColor(32, 76, 114))
            elif kind == "addon":
                _add_text(value, bold=True, size=9, color=RGBColor(20, 33, 52))
            elif kind == "channel":
                _add_text(value, bold=True, size=9)
            elif kind == "li":
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(value)
                run.font.name = "Arial"
                run.font.size = Pt(9.5)
            elif kind == "result_text":
                _add_text(value, bold=True, size=11, color=RGBColor(15, 81, 63))
            elif kind == "p":
                _add_text(value)
        flush_rows()

    def _add_plain_text_result(raw_text: str) -> None:
        report_html = _plain_result_to_report_html(raw_text, title)
        nodes = _report_nodes_from_html(report_html)
        if nodes:
            _add_report_nodes(nodes)
            return
        for line in raw_text.splitlines():
            if line.strip():
                _add_text(line.strip())
            else:
                doc.add_paragraph()

    def _extract_disclaimer(raw_text: str) -> str:
        marker = "Результат расчёта по настоящему калькулятору"
        pos = raw_text.find(marker)
        return raw_text[pos:].strip() if pos >= 0 else ""

    doc.add_heading("Результат расчёта", level=1)

    is_structured_radio_report = "radioReport" in report_html
    if req.summary and not is_structured_radio_report:
        rows = [(str(item.get("key", "")).strip(), str(item.get("value", "")).strip()) for item in req.summary]
        rows = [(k, v) for k, v in rows if k and v]
        if rows:
            _add_heading("Введённые данные", level=2)
            _add_rows_table([(k, v, False) for k, v in rows])

    if report_html:
        nodes = _report_nodes_from_html(report_html)
        if nodes:
            _add_report_nodes(nodes)
        else:
            _add_plain_text_result(text)
    else:
        _add_plain_text_result(text)

    disclaimer = _extract_disclaimer(text)
    if disclaimer:
        _add_text(disclaimer, italic=True, size=8, color=RGBColor(83, 101, 124))

    bio = io.BytesIO()
    doc.save(bio)
    content = bio.getvalue()
    headers = {
        "Content-Disposition": 'attachment; filename="rnk-calculation.docx"'
    }
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.post("/api/admin/rkn-upload")
@app.post("/api/admin/rkn-upload/")
async def api_admin_rkn_upload(
    request: Request,
    login: str = Form(...),
    password: str = Form(...),
    uploader: Optional[str] = Form(None),
    file: UploadFile = File(...),
):
    ok_auth, user = _check_admin_credentials(login, password)
    if not ok_auth:
        return JSONResponse(status_code=403, content={"ok": False, "error": "Неверный логин или пароль."})
    filename = str(file.filename or "").strip().lower()
    if not filename.endswith(".xlsx"):
        return JSONResponse(status_code=400, content={"ok": False, "error": "Нужен файл формата .xlsx."})

    content = await file.read()
    if not content:
        return JSONResponse(status_code=400, content={"ok": False, "error": "Файл пустой."})

    tmp_dir = PROJECT_DIR / ".tmp_rkn_upload_jobs"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    job_id = uuid.uuid4().hex
    raw_tmp = tmp_dir / f"{job_id}.xlsx"
    raw_tmp.write_bytes(content)
    _set_rkn_upload_job(
        job_id,
        status="queued",
        ok=True,
        message="Загрузка принята. Идёт обработка таблицы…",
        created_at=_now_msk().isoformat(timespec="seconds"),
        filename=str(file.filename or "").strip() or "uploaded.xlsx",
        uploader=str(uploader or "").strip() or None,
    )
    worker = threading.Thread(
        target=_run_rkn_upload_job,
        args=(
            job_id,
            raw_tmp,
            str(file.filename or "").strip() or "uploaded.xlsx",
            str(uploader or "").strip() or str((user or {}).get("name") or (user or {}).get("login") or ""),
            str(getattr(request.client, "host", "") or ""),
            str(request.headers.get("user-agent", "") or ""),
        ),
        daemon=True,
    )
    worker.start()
    return JSONResponse(
        status_code=202,
        content={
            "ok": True,
            "accepted": True,
            "job_id": job_id,
            "message": "Файл принят. Обработка выполняется в фоне, проверьте статус через job_id.",
        },
    )


@app.post("/api/admin/rkn-upload/auth")
@app.post("/api/admin/rkn-upload/auth/")
def api_admin_rkn_upload_auth(login: str = Form(...), password: str = Form(...)):
    ok_auth, user = _check_admin_credentials(login, password)
    if not ok_auth:
        return JSONResponse(status_code=403, content={"ok": False, "error": "Неверный логин или пароль."})
    return JSONResponse(
        status_code=200,
        content={
            "ok": True,
            "message": "Авторизация выполнена.",
            "user": {
                "login": str((user or {}).get("login") or ""),
                "name": str((user or {}).get("name") or ""),
            },
        },
    )


@app.get("/api/admin/rkn-upload/log")
@app.get("/api/admin/rkn-upload/log/")
def api_admin_rkn_upload_log(limit: int = 20):
    return {"ok": True, "log_path": str(RKN_UPLOAD_LOG_PATH), "items": _read_rkn_upload_log(limit=limit)}


@app.get("/api/admin/rkn-upload/status")
@app.get("/api/admin/rkn-upload/status/")
def api_admin_rkn_upload_status():
    artifacts = _rkn_artifacts_status()
    items = _read_rkn_upload_log(limit=1)
    latest = items[0] if items else _read_rkn_upload_state()
    def _ts(obj: Optional[dict[str, Any]]) -> Optional[datetime]:
        if not obj:
            return None
        raw = str(obj.get("timestamp") or "").strip()
        if not raw:
            return None
        try:
            return datetime.fromisoformat(raw.replace("Z", "+00:00"))
        except Exception:
            return None
    def _artifact_latest_ts() -> Optional[datetime]:
        cands = []
        for k in ("xlsx", "sqlite", "inn_csv"):
            raw = str((artifacts.get(k) or {}).get("mtime") or "").strip()
            if not raw:
                continue
            try:
                cands.append(datetime.fromisoformat(raw))
            except Exception:
                continue
        if not cands:
            return None
        return sorted(cands)[-1]

    if not latest:
        mtime_candidates = [
            artifacts.get("xlsx", {}).get("mtime"),
            artifacts.get("sqlite", {}).get("mtime"),
            artifacts.get("inn_csv", {}).get("mtime"),
        ]
        mtime_candidates = [x for x in mtime_candidates if x]
        if mtime_candidates:
            ts = sorted(mtime_candidates)[-1]
            latest = {
                "timestamp": ts,
                "filename": _registry_title_by_ts(ts),
                "uploader": None,
                "source": "server_artifact",
            }
    else:
        latest_ts = _ts(latest)
        art_ts = _artifact_latest_ts()
        if art_ts is not None and (latest_ts is None or art_ts > latest_ts):
            latest = {
                "timestamp": art_ts.isoformat(timespec="seconds"),
                "filename": _registry_title_by_ts(art_ts.isoformat(timespec="seconds")),
                "uploader": latest.get("uploader"),
                "source": "server_artifact",
            }
    return {
        "ok": True,
        "latest": latest,
        "artifacts": artifacts,
        "log_path": str(RKN_UPLOAD_LOG_PATH),
    }


@app.get("/api/admin/rkn-upload/job/{job_id}")
@app.get("/api/admin/rkn-upload/job/{job_id}/")
def api_admin_rkn_upload_job_status(job_id: str):
    item = _get_rkn_upload_job(job_id)
    if not item:
        return JSONResponse(status_code=404, content={"ok": False, "error": "Задача не найдена."})
    return {"ok": True, "job": item}


@app.get("/api/admin/rkn-upload/current-file")
@app.get("/api/admin/rkn-upload/current-file/")
def api_admin_rkn_upload_current_file():
    try:
        path = find_rkn_xlsx()
    except Exception:
        path = None
    if not path or not path.exists():
        return JSONResponse(status_code=404, content={"ok": False, "error": "Текущая таблица не найдена."})
    return FileResponse(
        path=str(path),
        filename=f"{_registry_title_by_ts()}.xlsx",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.get("/api/admin/usage-stats")
@app.get("/api/admin/usage-stats/")
def api_admin_usage_stats():
    report = _usage_reconcile_from_logs_monotonic(force=False)
    payload = {"ok": True, **_read_usage_stats()}
    if int(payload.get("total") or 0) <= 0:
        # Самовосстановление при "нуле" в панели: повторно прогоняем reconcile и backfill.
        try:
            report = _usage_reconcile_from_logs_monotonic(force=True) or report
            _usage_backfill_from_logs(patterns=None, dry_run=False)
            payload = {"ok": True, **_read_usage_stats()}
        except Exception:
            pass
        payload["diagnostics"] = _usage_scan_diagnostics()
    if report:
        payload["reconcile"] = report
    return JSONResponse(
        status_code=200,
        content=payload,
        headers={
            "Cache-Control": "no-store, no-cache, max-age=0, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )


@app.post("/api/admin/usage-stats/backfill")
@app.post("/api/admin/usage-stats/backfill/")
def api_admin_usage_stats_backfill(req: UsageStatsBackfillRequest):
    ok_auth, _user = _check_admin_credentials(req.login, req.password)
    if not ok_auth:
        return JSONResponse(status_code=403, content={"ok": False, "error": "Неверный логин или пароль."})
    try:
        reconcile_report = _usage_reconcile_from_logs_monotonic(force=True)
        report = _usage_backfill_from_logs(
            patterns=[str(x).strip() for x in (req.log_globs or []) if str(x).strip()],
            dry_run=bool(req.dry_run),
        )
        return {"ok": True, "reconcile": reconcile_report, "report": report, "stats": _read_usage_stats()}
    except Exception as e:
        return JSONResponse(status_code=500, content={"ok": False, "error": f"Ошибка backfill: {e}"})


@app.get("/api/version")
def api_version():
    return {"app": "app_fix", "rao": "rao_fix2", "calculators": ["rao_tv", "rao_radio", "vois_radio", "rao_events", "vois_tv", "vois_events"]}
